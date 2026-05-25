#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
统一访问形式集成测试 — {%tc for %} + {%tr for %} 全部使用 i.字段名 点号语法

覆盖：
  1. tc-for 表头 + 数据行统一使用 {{ i.xxx }}，列数全表统一
  2. tr-for 数据行 混合静态文本 + {{ i.xxx }} 动态字段
  3. tc + tr 联合，数据全部驱动化
  4. 动态列数响应（2项目/3项目）
"""

import sys
import tempfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from src.render.generator import DocumentGenerator
from src.processing.table_preprocessors import (
    TcInheritancePreprocessor,
    TrInheritancePreprocessor,
)

TESTS_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = TESTS_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


def make_template(path, rows_data):
    doc = Document()
    tbl = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]), style="Table Grid")
    for ri, row_cells in enumerate(rows_data):
        for ci, text in enumerate(row_cells):
            tbl.rows[ri].cells[ci].paragraphs[0].text = text if text else ""
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# 1. tc-for 统一访问
# ---------------------------------------------------------------------------
class TestTcForUnifiedAccess:
    """tc-for 表头和数据行均用 {{ i.xxx }}"""

    def test_uniform_columns_2_projects(self):
        """2 项目: 全表列数统一"""
        ctx = {
            "months": [
                {"标签": "1月", "值": "100"},
                {"标签": "2月", "值": "200"},
            ],
        }
        tmpl = OUTPUT_DIR / "_u_tc.docx"
        make_template(tmpl, [
            ["产品", "{%tc for m in months %}", "{{ m.标签 }}", "{%tc endfor %}"],
            ["产品A", "{{ m.值 }}", "", ""],
        ])

        prep = TcInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        prep.process(str(tmpl), tmp)
        DocumentGenerator(tmp).render(ctx, str(OUTPUT_DIR / "_u_tc_out.docx"))
        Path(tmp).unlink(missing_ok=True)

        doc = Document(str(OUTPUT_DIR / "_u_tc_out.docx"))
        t = doc.tables[0]
        col_nums = [len(r.cells) for r in t.rows]
        assert len(set(col_nums)) == 1, f"列数不统一: {col_nums}"
        assert len(t.rows[0].cells) == 3  # 1固定 + 2动态
        assert t.rows[0].cells[1].text.strip() == "1月"
        assert t.rows[0].cells[2].text.strip() == "2月"
        assert t.rows[1].cells[1].text.strip() == "100"
        assert t.rows[1].cells[2].text.strip() == "200"

    def test_uniform_columns_3_projects(self):
        """3 项目: 全表列数统一"""
        ctx = {
            "months": [
                {"标签": "1月", "值": "100"},
                {"标签": "2月", "值": "200"},
                {"标签": "3月", "值": "300"},
            ],
        }
        tmpl = OUTPUT_DIR / "_u_tc3.docx"
        make_template(tmpl, [
            ["产品", "{%tc for m in months %}", "{{ m.标签 }}", "{%tc endfor %}"],
            ["产品A", "{{ m.值 }}", "", ""],
        ])

        prep = TcInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        prep.process(str(tmpl), tmp)
        DocumentGenerator(tmp).render(ctx, str(OUTPUT_DIR / "_u_tc3_out.docx"))
        Path(tmp).unlink(missing_ok=True)

        doc = Document(str(OUTPUT_DIR / "_u_tc3_out.docx"))
        t = doc.tables[0]
        col_nums = [len(r.cells) for r in t.rows]
        assert len(set(col_nums)) == 1
        assert len(t.rows[0].cells) == 4  # 1固定 + 3动态


# ---------------------------------------------------------------------------
# 2. tr-for 静态+动态混排
# ---------------------------------------------------------------------------
class TestTrForMixedContent:
    """tr-for 内混合静态文本和 {{ i.xxx }} 动态字段"""

    def test_static_and_dynamic_coexist(self):
        """固定列 + 动态列 同时被循环"""
        ctx = {
            "项目列表": [
                {"项目": "项目X", "状态": "已完成"},
                {"项目": "项目Y", "状态": "进行中"},
            ],
        }
        tmpl = OUTPUT_DIR / "_u_tr.docx"
        make_template(tmpl, [
            ["序号", "项目", "状态"],
            ["{%tr for i in 项目列表 %}", "", ""],
            ["1", "{{ i.项目 }}", "{{ i.状态 }}"],
            ["{%tr endfor %}", "", ""],
        ])

        DocumentGenerator(str(tmpl)).render(ctx, str(OUTPUT_DIR / "_u_tr_out.docx"))

        doc = Document(str(OUTPUT_DIR / "_u_tr_out.docx"))
        t = doc.tables[0]
        assert len(t.rows) == 3
        assert t.rows[1].cells[0].text.strip() == "1"
        assert t.rows[1].cells[1].text.strip() == "项目X"
        assert t.rows[1].cells[2].text.strip() == "已完成"
        assert t.rows[2].cells[0].text.strip() == "1"
        assert t.rows[2].cells[1].text.strip() == "项目Y"
        assert t.rows[2].cells[2].text.strip() == "进行中"

    def test_fixed_text_outside_loop(self):
        """{%tr endfor %} 之后的固定行不受循环影响"""
        ctx = {"列表": [{"金额": 100}, {"金额": 200}]}
        tmpl = OUTPUT_DIR / "_u_tr_fixed.docx"
        make_template(tmpl, [
            ["金额"],
            ["{%tr for i in 列表 %}"],
            ["{{ i.金额 | money }}"],
            ["{%tr endfor %}"],
            ["合计"],
        ])

        DocumentGenerator(str(tmpl)).render(ctx, str(OUTPUT_DIR / "_u_tr_fixed_out.docx"))

        doc = Document(str(OUTPUT_DIR / "_u_tr_fixed_out.docx"))
        t = doc.tables[0]
        assert len(t.rows) == 4
        assert t.rows[-1].cells[0].text.strip() == "合计"


# ---------------------------------------------------------------------------
# 3. tc + tr 联合：全部数据驱动
# ---------------------------------------------------------------------------
class TestTcTrFullUnified:
    """tc + tr 联合，全部使用 i.xxx / m.xxx"""

    def test_full_unified_access(self):
        """二维动态表：行循环 × 列循环，全部数据驱动"""
        ctx = {
            "项目列表": [
                {"名称": "项目X", "审批": "已完成", "建设": "施工中"},
                {"名称": "项目Y", "审批": "审批中", "建设": "已完成"},
            ],
        }
        tmpl = OUTPUT_DIR / "_u_full.docx"
        make_template(tmpl, [
            ["项目", "{%tc for m in 阶段 %}", "{{ m.标签 }}", "{%tc endfor %}"],
            ["{%tr for i in 项目列表 %}", "", "", ""],
            ["{{ i.名称 }}", "{{ i[m.字段] }}", "", ""],
            ["{%tr endfor %}", "", "", ""],
        ])

        ctx["阶段"] = [
            {"标签": "审批", "字段": "审批"},
            {"标签": "建设", "字段": "建设"},
        ]

        prep = TcInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        prep.process(str(tmpl), tmp)
        DocumentGenerator(tmp).render(ctx, str(OUTPUT_DIR / "_u_full_out.docx"))
        Path(tmp).unlink(missing_ok=True)

        doc = Document(str(OUTPUT_DIR / "_u_full_out.docx"))
        t = doc.tables[0]
        col_nums = [len(r.cells) for r in t.rows]
        assert len(set(col_nums)) == 1, f"列数不统一: {col_nums}"
        assert t.rows[0].cells[1].text.strip() == "审批"
        assert t.rows[0].cells[2].text.strip() == "建设"
        assert t.rows[1].cells[0].text.strip() == "项目X"
        assert t.rows[1].cells[1].text.strip() == "已完成"
        assert t.rows[1].cells[2].text.strip() == "施工中"
        assert t.rows[2].cells[0].text.strip() == "项目Y"
        assert t.rows[2].cells[1].text.strip() == "审批中"
        assert t.rows[2].cells[2].text.strip() == "已完成"
