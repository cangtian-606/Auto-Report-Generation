#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试表格预处理器 — TcInheritancePreprocessor 和 TrInheritancePreprocessor

覆盖：
  Tc: 首行声明 {%tc for %} → 后续行自动注入列循环标签
  Tc: 已有 flow tags 的行不注入
  Tc: start cell 文本移动到 template cell
  Tr: 完整 tr-for/endfor 表不修改
  Tr: 缺少 {%tr endfor %} 时自动在表末插入新行补全
  Tr: 无 tr-for 表不修改
  Tc + Tr 串联不冲突
"""

import sys
import tempfile
import zipfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from lxml import etree
from src.processing.table_preprocessors import (
    TcInheritancePreprocessor,
    TrInheritancePreprocessor,
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TESTS_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = TESTS_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------
def read_template_cells(docx_path):
    """解压 docx 并读取所有表格的行列文本。"""
    tmpdir = tempfile.mkdtemp()
    with zipfile.ZipFile(docx_path) as z:
        z.extractall(tmpdir)
    tree = etree.parse(str(Path(tmpdir) / "word" / "document.xml"))
    result = []
    for tbl in tree.getroot().iter(f"{{{W}}}tbl"):
        table_rows = []
        for row in tbl.iter(f"{{{W}}}tr"):
            cells = []
            for tc in row.iter(f"{{{W}}}tc"):
                texts = [t.text or "" for t in tc.iter(f"{{{W}}}t")]
                cells.append("".join(texts))
            table_rows.append(cells)
        result.append(table_rows)
    return result


def make_clean_docx(path, rows_data):
    """创建干净的测试用 docx 模板。rows_data 每元素是 list of str。"""
    doc = Document()
    tbl = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]), style="Table Grid")
    for ri, row_cells in enumerate(rows_data):
        for ci, text in enumerate(row_cells):
            tbl.rows[ri].cells[ci].paragraphs[0].text = text if text else ""
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# TcInheritancePreprocessor 测试
# ---------------------------------------------------------------------------
class TestTcInheritancePreprocessor:
    """Tc: 首行声明 {%tc for %} → 后续行自动注入列循环标签"""

    def test_injects_tc_tags(self):
        """有 tc-for 首行时，后续行被注入 tc-for/endfor"""
        tmpl = OUTPUT_DIR / "_tc_inject.docx"
        make_clean_docx(tmpl, [
            ["固定", "{%tc for m in months %}", "{{ m }}", "{%tc endfor %}"],
            ["项目A", "是",             "",       ""],
        ])

        prep = TcInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        n = prep.process(str(tmpl), tmp)

        cells = read_template_cells(tmp)[0]
        Path(tmp).unlink(missing_ok=True)
        assert n == 1
        assert "{%tc for m in months  %}" in cells[1][1]
        assert "{%tc endfor %}" in cells[1][3]

    def test_preserves_start_cell_text(self):
        """start cell 的文本移动到 template cell，不丢失"""
        tmpl = OUTPUT_DIR / "_tc_preserve.docx"
        make_clean_docx(tmpl, [
            ["固定", "{%tc for m in months %}", "{{ m }}", "{%tc endfor %}"],
            ["项目A", "审批状态",             "",       ""],
        ])

        prep = TcInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        prep.process(str(tmpl), tmp)

        cells = read_template_cells(tmp)[0]
        Path(tmp).unlink(missing_ok=True)
        assert "审批状态" in cells[1][2]

    def test_skips_rows_with_flow_tags(self):
        """已有 tc/tr 标签的行不注入"""
        tmpl = OUTPUT_DIR / "_tc_skip.docx"
        make_clean_docx(tmpl, [
            ["固定", "{%tc for m in months %}", "{{ m }}", "{%tc endfor %}"],
            ["{%tr for i in 列表 %}", "", "", ""],
            ["{{ i.项目 }}", "值", "", ""],
        ])

        prep = TcInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        n = prep.process(str(tmpl), tmp)

        cells = read_template_cells(tmp)[0]
        Path(tmp).unlink(missing_ok=True)
        assert n == 1
        assert "{%tc for" not in cells[1][0]

    def test_no_tc_for_returns_zero(self):
        """无 tc-for 声明时返回 0"""
        tmpl = OUTPUT_DIR / "_tc_none.docx"
        make_clean_docx(tmpl, [
            ["项目", "金额", "状态"],
            ["项目A", "100", "已完成"],
        ])

        prep = TcInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        n = prep.process(str(tmpl), tmp)

        Path(tmp).unlink(missing_ok=True)
        assert n == 0

    def test_multi_row_injection(self):
        """多行数据全部注入"""
        tmpl = OUTPUT_DIR / "_tc_multi.docx"
        make_clean_docx(tmpl, [
            ["名称", "{%tc for m in months %}", "{{ m }}", "{%tc endfor %}"],
            ["行1", "是", "", ""],
            ["行2", "否", "", ""],
            ["行3", "是", "", ""],
        ])

        prep = TcInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        n = prep.process(str(tmpl), tmp)

        cells = read_template_cells(tmp)[0]
        Path(tmp).unlink(missing_ok=True)
        assert n == 1
        for ri in [1, 2, 3]:
            assert "{%tc for" in cells[ri][1]
            assert "{%tc endfor %}" in cells[ri][3]

    def test_single_cell_tc_for_start_end(self):
        """tc-for 和 tc-endfor 在相邻 cell (无独立 template cell) 时正常注入"""
        tmpl = OUTPUT_DIR / "_tc_single_cell.docx"
        make_clean_docx(tmpl, [
            ["名称", "{%tc for m in months %}", "{%tc endfor %}"],
            ["行1", "值", ""],
        ])

        prep = TcInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        n = prep.process(str(tmpl), tmp)

        cells = read_template_cells(tmp)[0]
        Path(tmp).unlink(missing_ok=True)
        assert n == 1
        assert "{%tc for" in cells[1][1]
        assert "{%tc endfor %}" in cells[1][2]


# ---------------------------------------------------------------------------
# TrInheritancePreprocessor 测试
# ---------------------------------------------------------------------------
class TestTrInheritancePreprocessor:
    """Tr: 检测 {%tr for %} 缺少 {%tr endfor %} 时自动补全"""

    def test_complete_tr_for_endfor_no_change(self):
        """完整 tr-for/endfor 表不修改"""
        tmpl = OUTPUT_DIR / "_tr_complete.docx"
        make_clean_docx(tmpl, [
            ["项目", "金额"],
            ["{%tr for i in 列表 %}", ""],
            ["{{ i.项目 }}", "{{ i.金额 }}"],
            ["{%tr endfor %}", ""],
        ])

        prep = TrInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        n = prep.process(str(tmpl), tmp)

        cells = read_template_cells(tmp)[0]
        Path(tmp).unlink(missing_ok=True)
        assert n == 0
        assert len(cells) == 4

    def test_missing_endfor_appends_new_row(self):
        """缺少 {%tr endfor %} 时在表末插入新行"""
        tmpl = OUTPUT_DIR / "_tr_missing.docx"
        make_clean_docx(tmpl, [
            ["项目", "金额"],
            ["{%tr for i in 列表 %}", ""],
            ["{{ i.项目 }}", "{{ i.金额 }}"],
        ])

        prep = TrInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        n = prep.process(str(tmpl), tmp)

        cells = read_template_cells(tmp)[0]
        Path(tmp).unlink(missing_ok=True)
        assert n == 1
        assert len(cells) == 4
        assert "{%tr endfor %}" in cells[3][0]

    def test_no_tr_for_returns_zero(self):
        """无 tr-for 表返回 0"""
        tmpl = OUTPUT_DIR / "_tr_none.docx"
        make_clean_docx(tmpl, [
            ["项目", "金额"],
            ["项目A", "100"],
        ])

        prep = TrInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        n = prep.process(str(tmpl), tmp)

        Path(tmp).unlink(missing_ok=True)
        assert n == 0

    def test_missing_endfor_with_tc_for(self):
        """有 tc-for + tr-for 无 endfor 时，Tc 不受影响"""
        tmpl = OUTPUT_DIR / "_tr_tc_missing.docx"
        make_clean_docx(tmpl, [
            ["名称", "{%tc for m in months %}", "{{ m }}", "{%tc endfor %}"],
            ["{%tr for i in 列表 %}", "", "", ""],
            ["{{ i.项目 }}", "是", "", ""],
        ])

        prep = TrInheritancePreprocessor()
        tmp = tempfile.mktemp(suffix=".docx")
        n = prep.process(str(tmpl), tmp)

        cells = read_template_cells(tmp)[0]
        Path(tmp).unlink(missing_ok=True)
        assert n == 1
        assert "{%tr endfor %}" in cells[-1][0]
        assert len(cells) == 4


# ---------------------------------------------------------------------------
# Tc + Tr 串联测试
# ---------------------------------------------------------------------------
class TestTcTrCombined:
    """Tc 和 Tr 顺序执行，不互相干扰"""

    def test_tc_after_tr_does_not_break(self):
        """先 Tr 补全，再 Tc 注入 — 两者不冲突"""
        tmpl = OUTPUT_DIR / "_combo.docx"
        make_clean_docx(tmpl, [
            ["名称", "{%tc for m in months %}", "{{ m }}", "{%tc endfor %}"],
            ["{%tr for i in 列表 %}", "", "", ""],
            ["{{ i.项目 }}", "是", "", ""],
        ])

        prep_tr = TrInheritancePreprocessor()
        prep_tc = TcInheritancePreprocessor()

        tmp1 = tempfile.mktemp(suffix=".docx")
        tmp2 = tempfile.mktemp(suffix=".docx")
        n_tr = prep_tr.process(str(tmpl), tmp1)
        n_tc = prep_tc.process(tmp1, tmp2)

        cells = read_template_cells(tmp2)[0]
        Path(tmp1).unlink(missing_ok=True)
        Path(tmp2).unlink(missing_ok=True)

        assert n_tr == 1
        assert n_tc == 1
        assert len(cells) == 4
        assert "{%tr endfor %}" in cells[3][0]
        assert "{%tc for" in cells[2][1]
        assert "{%tc endfor %}" in cells[2][3]
