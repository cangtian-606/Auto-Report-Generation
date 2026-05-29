"""pytest 批量处理测试"""

import os
import sys
from pathlib import Path

import pytest
from docx import Document

TESTS_DIR = Path(__file__).resolve().parent
PROJECT_DIR = TESTS_DIR.parent
sys.path.insert(0, str(PROJECT_DIR))

DATA_DIR = TESTS_DIR / "data"
TEMPLATE_DIR = TESTS_DIR / "templates"
TEMPLATE_PATH = TEMPLATE_DIR / "报告模板.docx"
OUTPUT_DIR = TESTS_DIR / "output" / "batch"


@pytest.fixture(scope="module")
def batch_setup():
    """准备2份测试数据 Excel，执行批量渲染"""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    if TEMPLATE_PATH.exists():
        TEMPLATE_PATH.unlink()

    _create_data_file(DATA_DIR / "data_公司A.xlsx", "深圳市A科技有限公司", "张A")
    _create_data_file(DATA_DIR / "data_公司B.xlsx", "北京市B有限公司", "李B")

    _setup_template()

    from src.cli import _render_batch
    results = _render_batch(
        str(DATA_DIR),
        str(TEMPLATE_PATH),
        str(OUTPUT_DIR),
        strict=False,
        check_vars=False,
    )
    results = [(f, s) for f, s in results if Path(f).stem.startswith("data_")]
    return results


def _create_data_file(path, company_name, legal_person):
    """创建测试数据 Excel"""
    import openpyxl

    wb = openpyxl.Workbook()

    ws_global = wb.active
    ws_global.title = "全局信息"
    ws_global.append(["字段编码", "值"])
    ws_global.append(["公司名", company_name])

    ws_basic = wb.create_sheet("基本情况")
    ws_basic.append(["字段编码", "值"])
    ws_basic.append(["法定代表人", legal_person])
    ws_basic.append(["注册资本", 500.00])
    ws_basic.append(["成立日期", "2020-01-01"])
    ws_basic.append(["有职工信息", "FALSE"])

    wb.save(str(path))


def _setup_template():
    """创建 docxtpl 模板（如果不存在）"""
    if TEMPLATE_PATH.exists():
        return

    from docx import Document

    doc = Document()
    doc.add_paragraph("公司名称：{{ 全局信息.公司名 }}")
    doc.add_paragraph("法定代表人：{{ 基本情况.法定代表人 }}")

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "金额"
    table.cell(0, 2).text = "比例"

    table.cell(1, 0).paragraphs[0].add_run("{%tr for i in 项目列表 %}")
    table.cell(2, 0).paragraphs[0].add_run("{{ i.项目 }}")
    table.cell(2, 1).paragraphs[0].add_run("{{ i.金额 | num }}")
    table.cell(2, 2).paragraphs[0].add_run("{{ i.比例 | percent }}")
    table.cell(3, 0).paragraphs[0].add_run("{%tr endfor %}")

    doc.save(str(TEMPLATE_PATH))


class TestBatchProcessing:
    def test_two_output_files_generated(self, batch_setup):
        results = batch_setup
        assert len(results) == 2, f"期望2个文件，实际{len(results)}个"

    def test_company_a_output_exists(self, batch_setup):
        assert (OUTPUT_DIR / "公司A_output.docx").exists()

    def test_company_b_output_exists(self, batch_setup):
        assert (OUTPUT_DIR / "公司B_output.docx").exists()

    def test_both_files_successful(self, batch_setup):
        results = batch_setup
        success_count = sum(1 for _, s in results if s)
        assert success_count == 2, f"期望2个成功，实际{success_count}个"

    def test_company_a_content_replaced(self, batch_setup):
        output_a = OUTPUT_DIR / "公司A_output.docx"
        if not output_a.exists():
            pytest.skip("输出文件不存在，跳过内容验证")
        doc = Document(str(output_a))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "深圳市A科技有限公司" in all_text

    def test_company_b_content_replaced(self, batch_setup):
        output_b = OUTPUT_DIR / "公司B_output.docx"
        if not output_b.exists():
            pytest.skip("输出文件不存在，跳过内容验证")
        doc = Document(str(output_b))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "北京市B有限公司" in all_text
