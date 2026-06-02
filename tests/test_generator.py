#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试文档生成器功能
"""

import os
import tempfile
import pytest
from pathlib import Path
from docx import Document

from src.render.generator import DocumentGenerator
from src.render.analyzer import TemplateAnalyzer


@pytest.fixture
def temp_dir():
    """临时目录 fixture"""
    with tempfile.TemporaryDirectory() as tmp:
        yield Path(tmp)


@pytest.fixture
def simple_template(temp_dir):
    """创建一个简单的测试模板"""
    tmpl_path = temp_dir / "test_template.docx"
    doc = Document()
    doc.add_paragraph("公司名称: {{ global.company_name }}")
    doc.add_paragraph("成立日期: {{ global.found_date }}")
    doc.add_paragraph("项目列表:")
    # 简化为不使用 tr 循环，避免 docxtpl 特定语法问题
    doc.add_paragraph("示例项目: {{ projects[0].name if projects else '' }}")
    doc.save(tmpl_path)
    return tmpl_path


class TestDocumentGenerator:
    """测试 DocumentGenerator 类"""

    def test_find_unused_data(self, simple_template):
        """测试未使用数据检测功能"""
        analyzer = TemplateAnalyzer(str(simple_template))
        ctx = {
            "global": {
                "company_name": "Test",
                "found_date": "2020-01-01",
                "unused_field": "value"
            },
            "projects": [
                {"name": "Proj1", "year": 2020, "unused_proj_field": "v"}
            ],
            "unused_section": "value"
        }

        unused = analyzer.get_unused(ctx)

        # 检查未使用字段是否被识别
        assert "global.unused_field" in unused
        assert "unused_section" in unused

        # 已使用的字段不应该出现在未使用列表
        assert "global.company_name" not in unused
        assert "global.found_date" not in unused

    def test_render(self, simple_template, temp_dir):
        """测试基本渲染功能"""
        gen = DocumentGenerator(str(simple_template))
        output_path = temp_dir / "output.docx"
        ctx = {
            "global": {
                "company_name": "Test Company",
                "found_date": "2020-01-01"
            },
            "projects": [
                {"name": "Project 1", "year": 2020},
                {"name": "Project 2", "year": 2021}
            ]
        }

        result = gen.render(ctx, str(output_path))

        assert result is True
        assert output_path.exists()
        stat = output_path.stat()
        assert stat.st_size > 0

    def test_check_syntax_no_errors(self, simple_template):
        """测试变量检查在无错误情况下"""
        analyzer = TemplateAnalyzer(str(simple_template))
        ctx = {
            "global": {
                "company_name": "Test",
                "found_date": "2020-01-01"
            },
            "projects": [{"name": "P1", "year": 2020}]
        }

        undeclared = analyzer.get_undeclared(ctx)
        assert undeclared is not None
        assert len(undeclared) == 0

    def test_for_loop_iterated_list_fields_not_reported(self, temp_dir):
        """模板中 {% for %} 块迭代的列表，其内部字段不应被误报为未用。

        验证 get_unused 不会生成 projects[0].name 这样的路径。
        """
        from docx import Document
        tmpl_path = temp_dir / "for_template.docx"
        doc = Document()
        doc.add_paragraph(
            "{% for item in projects %}{{ item.name }} - {{ item.year }}{% endfor %}"
        )
        doc.save(tmpl_path)

        analyzer = TemplateAnalyzer(str(tmpl_path))
        ctx = {
            "projects": [
                {"name": "Proj1", "year": 2020, "secret": "x"},
                {"name": "Proj2", "year": 2021, "secret": "y"},
            ]
        }

        unused = analyzer.get_unused(ctx)

        # 整个 projects 列表被 for 块引用，不应被报告
        assert "projects" not in unused
        # 列表内任何字段都不应被下钻报告（不下钻 list 内部）
        assert not any(p.startswith("projects[") for p in unused), (
            f"不应下钻 list 内部: {unused}"
        )

    def test_completely_unused_list_is_reported(self, temp_dir):
        """完全未被引用的 list 整体应被报告为未用。"""
        from docx import Document
        tmpl_path = temp_dir / "unused_list_template.docx"
        doc = Document()
        doc.add_paragraph("公司: {{ company.name }}")
        doc.save(tmpl_path)

        analyzer = TemplateAnalyzer(str(tmpl_path))
        ctx = {
            "company": {"name": "Acme"},
            "orphan_projects": [{"name": "P1"}, {"name": "P2"}],
        }

        unused = analyzer.get_unused(ctx)

        assert "orphan_projects" in unused
        # company.name 仍正常识别为已用
        assert "company.name" not in unused
        # orphan_projects 内字段不应被下钻
        assert not any(p.startswith("orphan_projects[") for p in unused), (
            f"不应下钻 list 内部: {unused}"
        )

    def test_template_cache(self, simple_template):
        """测试模板缓存机制"""
        gen1 = DocumentGenerator(str(simple_template))
        doc1 = gen1._load_template()
        doc2 = gen1._load_template()

        # 同一对象引用
        assert doc1 is doc2

        # 不同实例相同路径，使用同一缓存
        gen2 = DocumentGenerator(str(simple_template))
        doc3 = gen2._load_template()
        assert doc3 is doc1

    def test_template_cache_invalidates_on_mtime_change(self, temp_dir):
        """模板文件 mtime 变化时缓存应自动失效（重新加载）"""
        from docx import Document
        tmpl_path = temp_dir / "mtime_template.docx"
        Document().save(tmpl_path)

        gen = DocumentGenerator(str(tmpl_path))
        doc1 = gen._load_template()

        # 修改 mtime：后退 10 秒，确保 mtime_ns 必变
        import os
        old_mtime = os.stat(tmpl_path).st_mtime
        os.utime(tmpl_path, (old_mtime - 10, old_mtime - 10))

        doc2 = gen._load_template()

        # mtime 变化后，缓存应返回新实例
        assert doc1 is not doc2, "mtime 变化后缓存应失效"

    def test_load_template_log_uses_source_path(self, temp_dir, caplog):
        """DEBUG 日志应显示原始模板路径（source_template_path），而非预处理后临时路径。

        场景：DocumentGenerator 接收的是预处理后临时文件，但日志需显示用户原始路径。
        """
        import logging
        from docx import Document

        source_path = temp_dir / "user_template.docx"
        preprocessed_path = temp_dir / "tc_pre_abc.docx"
        Document().save(source_path)
        Document().save(preprocessed_path)

        # 创建 generator 时传入预处理路径，但显式指定原始路径
        gen = DocumentGenerator(
            template_path=str(preprocessed_path),
            source_template_path=str(source_path),
        )

        with caplog.at_level(logging.DEBUG, logger="src.render.generator"):
            gen._load_template()

        # 日志中应该出现原始路径
        log_text = caplog.text
        assert str(source_path) in log_text, (
            f"日志应显示原始路径 {source_path}，实际: {log_text}"
        )

    def test_tc_column_loop(self, temp_dir):
        """测试 {%tc for %} 列循环渲染 — 表头 + 多行数据"""
        from docx import Document
        tmpl_path = temp_dir / "tc_template.docx"
        out_path = temp_dir / "tc_output.docx"

        doc = Document()
        table = doc.add_table(rows=3, cols=5, style="Table Grid")

        # Row 0: 表头 — tc 循环动态生成月份列
        h = table.rows[0]
        h.cells[0].text = "产品"
        h.cells[1].paragraphs[0].text = "{%tc for month in months %}"
        h.cells[2].paragraphs[0].text = "{{ month }}"
        h.cells[3].paragraphs[0].text = "{%tc endfor %}"
        h.cells[4].text = "合计"

        # Row 1: 数据行 A
        a = table.rows[1]
        a.cells[0].text = "产品A"
        a.cells[1].paragraphs[0].text = "{%tc for month in months %}"
        a.cells[2].paragraphs[0].text = "{{ data_a[month] | num }}"
        a.cells[3].paragraphs[0].text = "{%tc endfor %}"
        a.cells[4].text = "{{ total_a | num }}"

        # Row 2: 数据行 B
        b = table.rows[2]
        b.cells[0].text = "产品B"
        b.cells[1].paragraphs[0].text = "{%tc for month in months %}"
        b.cells[2].paragraphs[0].text = "{{ data_b[month] | num }}"
        b.cells[3].paragraphs[0].text = "{%tc endfor %}"
        b.cells[4].text = "{{ total_b | num }}"

        doc.save(tmpl_path)

        gen = DocumentGenerator(str(tmpl_path))
        ctx = {
            "months": ["1月", "2月", "3月"],
            "data_a": {"1月": 100, "2月": 200, "3月": 150},
            "total_a": 450,
            "data_b": {"1月": 80, "2月": 120, "3月": 90},
            "total_b": 290,
        }

        result = gen.render(ctx, str(out_path))
        assert result is True
        assert out_path.exists()

        out_doc = Document(str(out_path))
        cells = [c.text.strip() for c in out_doc.tables[0].rows[0].cells]
        assert cells[0] == "产品"
        assert cells[1] == "1月"
        assert cells[2] == "2月"
        assert cells[3] == "3月"
        assert cells[4] == "合计"

        row_a = [c.text.strip() for c in out_doc.tables[0].rows[1].cells]
        assert row_a[0] == "产品A"
        assert row_a[1] == "100.00"
        assert row_a[4] == "450.00"

    def test_docx_editor_exit_sets_tmp_dir_none(self, temp_dir):
        """DocxEditor __exit__ 后应将 _tmp_dir 设为 None，且不能在块外调用 save_to。"""
        from docx import Document
        from src.processing.docx_editor import DocxEditor

        tmpl_path = temp_dir / "test_editor.docx"
        out1_path = temp_dir / "out1.docx"
        out2_path = temp_dir / "out2.docx"
        doc = Document()
        doc.add_paragraph("test")
        doc.save(tmpl_path)

        # 在 with 块内正常工作
        editor = DocxEditor(str(tmpl_path))
        with editor:
            # 检查内部状态
            assert editor._tmp_dir is not None
            editor.save_to(str(out1_path), modified=True)

        # 检查 with 块退出后的状态
        assert editor._tmp_dir is None
        # 尝试在块外调用 save_to 应该抛 AssertionError
        try:
            editor.save_to(str(out2_path), modified=False)
            assert False, "应在退出后抛 AssertionError"
        except AssertionError as e:
            assert "只能在 with 块内" in str(e)
