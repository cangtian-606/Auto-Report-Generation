import os
import sys
from pathlib import Path

PROJECT_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_DIR))

from docx import Document

TEMPLATE_SRC = PROJECT_DIR / "templates" / "报告测试.docx"
TEMPLATE_OUT = PROJECT_DIR / "templates" / "报告模板.docx"
DATA_FILE = PROJECT_DIR / "data" / "报告数据.xlsx"
OUTPUT_FILE = PROJECT_DIR / "output" / "投资项目概况_测试输出.docx"
EXIT_SUCCESS = 0
EXIT_FAILURE = 1


def setup_template():
    """将原始报告测试.docx转换为含Jinja2变量的docxtpl模板"""
    from docx import Document

    doc = Document(str(TEMPLATE_SRC))

    # === 表格0: 基本信息 ===
    basic_fields = [
        ("date.全局信息.公司名", None),
        ("date.基本情况.信用代码", None),
        ("date.基本情况.注册地址", None),
        ("date.基本情况.法定代表人", None),
        ("date.基本情况.公司类型", None),
        ("date.基本情况.注册资本", "num"),
        ("date.基本情况.经营范围", None),
        ("date.基本情况.成立日期", None),
        ("date.基本情况.经营期限", None),
    ]
    table0 = doc.tables[0]
    for ri, (field, flt) in enumerate(basic_fields):
        cell = table0.rows[ri].cells[1]
        para = cell.paragraphs[0]
        para.clear()
        if flt:
            para.add_run("{{ " + field + " | " + flt + " }}")
        else:
            para.add_run("{{ " + field + " }}")

    # === 表格1: 股东出资（动态循环）===
    table1 = doc.tables[1]
    # 清空空行，只保留表头行
    while len(table1.rows) > 1:
        tbl_elem = table1.rows[-1]._tr.getparent()
        tbl_elem.remove(table1.rows[-1]._tr)
    # 添加 for 行
    for_row1 = table1.add_row()
    for_row1.cells[0].paragraphs[0].add_run("{%tr for i in form.股东出资 %}")
    for_row1.cells[1].paragraphs[0].clear()
    for_row1.cells[2].paragraphs[0].clear()
    for_row1.cells[3].paragraphs[0].clear()
    for_row1.cells[4].paragraphs[0].clear()
    # 添加数据行（认缴/实缴金额用money过滤器，比例用percent过滤器）
    data_row1 = table1.add_row()
    data_row1.cells[0].paragraphs[0].add_run("{{ i.股东 }}")
    data_row1.cells[1].paragraphs[0].add_run("{{ i.认缴金额 | money }}")
    data_row1.cells[2].paragraphs[0].add_run("{{ i.认缴比例 | percent }}")
    data_row1.cells[3].paragraphs[0].add_run("{{ i.实缴金额 | money }}")
    data_row1.cells[4].paragraphs[0].add_run("{{ i.实缴比例 | percent }}")
    # 添加 endfor 行
    endfor_row1 = table1.add_row()
    endfor_row1.cells[0].paragraphs[0].add_run("{%tr endfor %}")

    # === 表格2: 员工信息（动态循环，有条件）===
    table2 = doc.tables[2]
    table3 = doc.tables[3]
    if_para = doc.add_paragraph("{%p if date.基本情况.有职工信息 %}")
    end_para = doc.add_paragraph("{%p endif %}")
    body = doc.element.body
    body.insert(list(body).index(table2._element), if_para._p)
    body.insert(list(body).index(table3._element), end_para._p)

    while len(table2.rows) > 1:
        tbl_elem = table2.rows[1]._tr.getparent()
        tbl_elem.remove(table2.rows[1]._tr)
    for_row2 = table2.add_row()
    for_row2.cells[0].paragraphs[0].add_run("{%tr for i in form.员工信息 %}")
    data_row2 = table2.add_row()
    data_row2.cells[0].paragraphs[0].add_run("{{ i.序号 }}")
    data_row2.cells[1].paragraphs[0].add_run("{{ i.职务 }}")
    endfor_row2 = table2.add_row()
    endfor_row2.cells[0].paragraphs[0].add_run("{%tr endfor %}")

    # === 表格3: 对外投资明细表（动态循环）===
    table3 = doc.tables[3]
    while len(table3.rows) > 1:
        tbl_elem = table3.rows[1]._tr.getparent()
        tbl_elem.remove(table3.rows[1]._tr)
    for_row3 = table3.add_row()
    for_row3.cells[0].paragraphs[0].add_run("{%tr for i in form.投资明细 %}")
    data_row3 = table3.add_row()
    data_row3.cells[0].paragraphs[0].add_run("{{ i.项目 }}")
    data_row3.cells[1].paragraphs[0].add_run("{{ i.金额 | money }}")
    data_row3.cells[2].paragraphs[0].add_run("{{ i.持股比例 | percent }}")
    endfor_row3 = table3.add_row()
    endfor_row3.cells[0].paragraphs[0].add_run("{%tr endfor %}")

    doc.save(str(TEMPLATE_OUT))
    return True


def setup_data():
    """创建测试数据Excel"""
    import openpyxl
    from openpyxl.styles import Font

    wb = openpyxl.Workbook()

    # Sheet 1: date.全局信息
    ws_global = wb.active
    ws_global.title = "date.全局信息"
    ws_global.append(["字段编码", "值"])
    global_data = [
        ["公司名", "深圳市创新科技有限公司"],
    ]
    for row_data in global_data:
        ws_global.append(row_data)
    ws_global.column_dimensions["A"].width = 25
    ws_global.column_dimensions["B"].width = 50

    # Sheet 2: date.基本情况
    ws_basic = wb.create_sheet("date.基本情况")
    ws_basic.append(["字段编码", "值"])
    basic_data = [
        ["信用代码", "91440300MA5XXXXX00"],
        ["注册地址", "深圳市南山区粤海街道科技园南路88号"],
        ["法定代表人", "张伟"],
        ["公司类型", "有限责任公司"],
        ["注册资本", 1000.00],
        ["经营范围", "软件开发、技术咨询、技术服务"],
        ["成立日期", "2018-06-15"],
        ["经营期限", "2018-06-15 至 长期"],
        ["有职工信息", "FALSE"],
    ]
    for row_data in basic_data:
        ws_basic.append(row_data)
    ws_basic.column_dimensions["A"].width = 25
    ws_basic.column_dimensions["B"].width = 50

    # Sheet 3: form.股东出资
    ws_sh = wb.create_sheet("form.股东出资")
    ws_sh.append([
        "股东",
        "认缴金额",
        "认缴比例",
        "实缴金额",
        "实缴比例",
    ])
    shareholders = [
        ["张伟", 600.00, 0.60, 600.00, 0.60],
        ["李芳", 250.00, 0.25, 250.00, 0.25],
        ["王强", 150.00, 0.15, 150.00, 0.15],
    ]
    for row_data in shareholders:
        ws_sh.append(row_data)
    ws_sh.column_dimensions["A"].width = 20
    for col in ["B", "C", "D", "E"]:
        ws_sh.column_dimensions[col].width = 18

    # Sheet 4: form.员工信息
    ws_info = wb.create_sheet("form.员工信息")
    ws_info.append(["序号", "职务"])
    info_data = [
        [1, "张伟-执行董事"],
        [2, "李芳-监事"],
        [3, "王强-股东"],
    ]
    for row_data in info_data:
        ws_info.append(row_data)
    ws_info.column_dimensions["A"].width = 10
    ws_info.column_dimensions["B"].width = 25

    # Sheet 5: form.投资明细
    ws_invest = wb.create_sheet("form.投资明细")
    ws_invest.append(["项目", "金额", "持股比例"])
    invest_data = [
        ["项目A", 500.00, 0.50],
        ["项目B", 300.00, 0.30],
        ["项目C", 200.00, 0.20],
        ["合计", 1000.00, 1.00],
    ]
    for row_data in invest_data:
        ws_invest.append(row_data)
    ws_invest.column_dimensions["A"].width = 20
    ws_invest.column_dimensions["B"].width = 18
    ws_invest.column_dimensions["C"].width = 15

    wb.save(str(DATA_FILE))
    return True


def run_render():
    """调用 generator 执行渲染"""
    from src.generator import generate

    success = generate(
        str(DATA_FILE),
        str(TEMPLATE_OUT),
        str(OUTPUT_FILE),
        strict=False,
        check_vars=True,
    )
    return success


def verify_output():
    """验证渲染输出文档内容正确"""
    doc = Document(str(OUTPUT_FILE))
    all_text = " ".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                all_text += " " + c.text

    checks = {
        "公司名称已替换": "深圳市创新科技有限公司" in all_text,
        "信用代码已替换": "91440300MA5XXXXX00" in all_text,
        "法定代表人已替换": "张伟" in all_text,
        "注册资本已替换": "1,000" in all_text,
        "成立日期已替换": "2018-06-15" in all_text,
        "股东李芳在表格中": "李芳" in all_text,
        "认缴金额已替换": "250" in all_text,
        "金额格式_认缴250": "250.00" in all_text or "250,00" in all_text,
        "金额格式_实缴600": "600.00" in all_text or "600,00" in all_text,
        "百分比格式_25%": "25.00%" in all_text,
        "百分比格式_60%": "60.00%" in all_text,
        "项目A在投资明细中": "项目A" in all_text,
        "合计在投资明细中": "合计" in all_text,
        "投资金额500格式": "500.00" in all_text or "500,00" in all_text,
        "持股比例50%格式": "50.00%" in all_text,
        "有职工信息FALSE时隐藏职工表格": "张伟-执行董事" not in all_text,
        "无残留Jinja2标签": "{{" not in all_text and "{%" not in all_text,
    }

    all_pass = True
    print("\n" + "=" * 50)
    print("输出验证:")
    for name, result in checks.items():
        status = "PASS" if result else "FAIL"
        print("  [{}] {}".format(status, name))
        if not result:
            all_pass = False
    print("=" * 50)

    return all_pass


def main():
    print("=" * 50)
    print("TDD: 报告模板渲染测试")
    print("=" * 50)

    # Step 1: 创建模板 (从原始报告测试.docx转换)
    print("\n[SETUP] 创建模板...")
    if not setup_template():
        print("ERROR: 模板创建失败")
        return EXIT_FAILURE
    print("  模板已生成: {}".format(TEMPLATE_OUT))

    # Step 2: 创建测试数据
    print("\n[SETUP] 创建测试数据...")
    if not setup_data():
        print("ERROR: 数据创建失败")
        return EXIT_FAILURE
    print("  数据已生成: {}".format(DATA_FILE))

    # Step 3: 执行渲染
    print("\n[RENDER] 执行渲染...")
    os.makedirs(OUTPUT_FILE.parent, exist_ok=True)
    success = run_render()
    if not success:
        print("ERROR: 渲染失败")
        return EXIT_FAILURE
    print("  输出已生成: {}".format(OUTPUT_FILE))

    # Step 4: 验证输出
    all_pass = verify_output()

    if all_pass:
        print("\n✓ 所有检查通过 - 测试通过!")
        return EXIT_SUCCESS
    else:
        print("\n✗ 存在失败检查 - 测试未通过")
        return EXIT_FAILURE


if __name__ == "__main__":
    exit_code = main()
    sys.exit(exit_code)
