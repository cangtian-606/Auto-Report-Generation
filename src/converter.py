#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word模板转换器（通用版）

将Word文档转换为docxtpl模板，支持：
- 占位符识别与替换
- 全局变量替换
- 操作指引清理
- 页眉页脚处理

使用方法：
    python word_template_converter.py input.docx output.docx
    python word_template_converter.py --config config.yaml input.docx output.docx

作者：通用模板工具
版本：1.0.0
"""

import os
import re
import sys
import json
import argparse
import logging
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Callable, Any
from dataclasses import dataclass, asdict

try:
    from docx import Document
except ImportError:
    print("错误：请先安装 python-docx: pip install python-docx")
    sys.exit(1)


# ============================================================
# 配置类
# ============================================================

@dataclass
class ReplacementConfig:
    """替换配置"""
    # 全局变量替换映射 [(原始文本, 替换为), ...]
    global_replacements: List[Tuple[str, str]] = None
    
    # 占位符正则模式
    placeholder_pattern: str = r"(?<![a-zA-Z_])[×X]{4,}(?![a-zA-Z_}])"
    
    # 占位符替换模板，{desc} 会被替换为描述标识
    placeholder_template: str = "{{ notes.placeholder_{desc} }}"
    
    # 操作指引清理正则
    note_pattern: str = r"【注[：:].*?】?"
    
    # 是否启用占位符描述推断
    infer_descriptions: bool = True
    
    # 描述推断的上下文长度（字符数）
    context_length: int = 60
    
    def __post_init__(self):
        if self.global_replacements is None:
            self.global_replacements = []


@dataclass
class ProcessingStats:
    """处理统计"""
    paragraphs_processed: int = 0
    tables_processed: int = 0
    headers_footers_processed: int = 0
    placeholders_replaced: int = 0
    globals_replaced: int = 0
    notes_removed: int = 0
    
    def to_dict(self) -> Dict[str, int]:
        return asdict(self)


# ============================================================
# 日志配置
# ============================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)],
)
logger = logging.getLogger(__name__)


# ============================================================
# 占位符描述生成器
# ============================================================

class DescriptionGenerator:
    """
    占位符描述标识生成器
    
    根据上下文提取关键词生成有意义的变量名。
    """
    
    def __init__(self):
        self.counter = 0
        self.used_names: set = set()
    
    def next_id(self) -> str:
        """生成下一个序号标识"""
        self.counter += 1
        return f"{self.counter:04d}"
    
    def infer_description(self, full_text: str, match_start: int, 
                         context_length: int = 60) -> Optional[str]:
        """
        从上下文推断占位符描述
        
        Args:
            full_text: 段落完整文本
            match_start: 匹配起始位置
            context_length: 上下文长度
            
        Returns:
            推断的描述标识，无法推断时返回None
        """
        # 取匹配位置前的文本作为上下文
        ctx_start = max(0, match_start - context_length)
        before = full_text[ctx_start:match_start].strip()
        
        desc = None
        
        # 模式1: "xxx：" 或 "xxx:" 后面紧跟占位符
        colon_match = re.search(r"[\uff1a:]([^：:\s]{2,20})\s*$", before)
        if colon_match:
            desc = colon_match.group(1).strip()
            desc = re.sub(r"^[（(]|[）)]$", "", desc)
        
        # 模式2: "的" 前面的名词
        if not desc:
            de_match = re.search(r"([\u4e00-\u9fff]{2,10})的\s*$", before)
            if de_match:
                desc = de_match.group(1)
        
        # 模式3: "为"或"是"前面的名词
        if not desc:
            verb_match = re.search(r"([\u4e00-\u9fff]{2,15})\s*(?:为|是)\s*$", before)
            if verb_match:
                desc = verb_match.group(1)
        
        if desc:
            # 清理desc，只保留字母、数字、中文、下划线
            desc = re.sub(r"[^\u4e00-\u9fff\w]", "_", desc)
            desc = re.sub(r"_+", "_", desc).strip("_")
            if len(desc) > 40:
                desc = desc[:40]
            if desc:
                # 确保唯一性
                original_desc = desc
                suffix = 1
                while desc in self.used_names:
                    desc = f"{original_desc}_{suffix}"
                    suffix += 1
                self.used_names.add(desc)
                return desc
        
        return None
    
    def get_unique_id(self) -> str:
        """获取唯一的序号标识"""
        desc = self.next_id()
        while desc in self.used_names:
            desc = self.next_id()
        self.used_names.add(desc)
        return desc


# ============================================================
# 文本替换器
# ============================================================

class TextReplacer:
    """文本替换器"""
    
    def __init__(self, config: ReplacementConfig):
        self.config = config
        self.desc_generator = DescriptionGenerator()
        self.placeholder_regex = re.compile(config.placeholder_pattern)
        self.note_regex = re.compile(config.note_pattern)
    
    def clean_notes(self, text: str) -> str:
        """清理操作指引"""
        return self.note_regex.sub("", text)
    
    def replace_globals(self, text: str) -> Tuple[str, int]:
        """执行全局变量替换"""
        count = 0
        for old, new in self.config.global_replacements:
            if old in text:
                text = text.replace(old, new)
                count += 1
        return text, count
    
    def replace_placeholders(self, text: str) -> Tuple[str, int]:
        """替换占位符"""
        count = 0
        
        def _replace_match(m: re.Match) -> str:
            nonlocal count
            count += 1
            
            if self.config.infer_descriptions:
                desc = self.desc_generator.infer_description(
                    text, m.start(), self.config.context_length
                )
                if not desc:
                    desc = self.desc_generator.get_unique_id()
            else:
                desc = self.desc_generator.get_unique_id()
            
            return self.config.placeholder_template.format(desc=desc)
        
        new_text = self.placeholder_regex.sub(_replace_match, text)
        return new_text, count
    
    def process(self, text: str) -> Tuple[str, Dict[str, int]]:
        """
        执行完整的替换流程
        
        Returns:
            (新文本, 统计信息)
        """
        stats = {"globals": 0, "placeholders": 0, "notes": 0}
        
        # 步骤1: 清理操作指引
        original_len = len(text)
        text = self.clean_notes(text)
        if len(text) < original_len:
            stats["notes"] = 1
        
        # 步骤2: 全局变量替换
        text, globals_count = self.replace_globals(text)
        stats["globals"] = globals_count
        
        # 步骤3: 占位符替换
        text, placeholder_count = self.replace_placeholders(text)
        stats["placeholders"] = placeholder_count
        
        return text, stats


# ============================================================
# 文档处理器
# ============================================================

class DocumentProcessor:
    """Word文档处理器"""
    
    def __init__(self, replacer: TextReplacer):
        self.replacer = replacer
        self.stats = ProcessingStats()
    
    def process_paragraph(self, paragraph) -> bool:
        """处理单个段落"""
        runs = paragraph.runs
        if not runs:
            return False
        
        # 合并所有Run文本
        full_text = "".join(run.text for run in runs)
        if not full_text:
            return False
        
        # 执行替换
        new_text, replace_stats = self.replacer.process(full_text)
        
        if new_text == full_text:
            return False
        
        # 写回第一个Run，清空其余
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
        
        # 更新统计
        self.stats.globals_replaced += replace_stats["globals"]
        self.stats.placeholders_replaced += replace_stats["placeholders"]
        self.stats.notes_removed += replace_stats["notes"]
        
        return True
    
    def process_table(self, table) -> int:
        """处理表格"""
        count = 0
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if self.process_paragraph(para):
                        count += 1
        return count
    
    def process_header_footer(self, header_or_footer) -> int:
        """处理页眉或页脚"""
        count = 0
        
        # 处理段落
        for para in header_or_footer.paragraphs:
            if self.process_paragraph(para):
                count += 1
        
        # 处理表格
        for table in header_or_footer.tables:
            count += self.process_table(table)
        
        return count
    
    def process(self, doc: Document) -> ProcessingStats:
        """处理整个文档"""
        # 处理正文段落
        for para in doc.paragraphs:
            if self.process_paragraph(para):
                self.stats.paragraphs_processed += 1
        
        # 处理表格
        for table in doc.tables:
            replaced = self.process_table(table)
            if replaced > 0:
                self.stats.tables_processed += replaced
        
        # 处理页眉页脚
        for section in doc.sections:
            if section.header:
                replaced = self.process_header_footer(section.header)
                if replaced > 0:
                    self.stats.headers_footers_processed += replaced
            
            if section.footer:
                replaced = self.process_header_footer(section.footer)
                if replaced > 0:
                    self.stats.headers_footers_processed += replaced
        
        return self.stats


# ============================================================
# 验证器
# ============================================================

class TemplateValidator:
    """模板验证器"""
    
    def __init__(self, placeholder_pattern: str = r"[×X]{4,}", 
                 note_pattern: str = r"【注"):
        self.placeholder_regex = re.compile(placeholder_pattern)
        self.note_marker = note_pattern
    
    def validate(self, doc: Document) -> Dict[str, Any]:
        """验证模板转换结果"""
        results = {
            "placeholder_residuals": 0,
            "note_residuals": 0,
            "jinja2_tags": 0,
            "unclosed_tags": 0,
            "issues": []
        }
        
        all_texts = []
        
        # 收集所有文本
        for para in doc.paragraphs:
            text = "".join(r.text for r in para.runs)
            all_texts.append(text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = "".join(r.text for r in para.runs)
                        all_texts.append(text)
        
        for section in doc.sections:
            for para in section.header.paragraphs:
                text = "".join(r.text for r in para.runs)
                all_texts.append(text)
            for para in section.footer.paragraphs:
                text = "".join(r.text for r in para.runs)
                all_texts.append(text)
        
        # 检查
        for text in all_texts:
            # 占位符残留
            if self.placeholder_regex.search(text):
                results["placeholder_residuals"] += 1
                if len(results["issues"]) < 5:
                    results["issues"].append(f"占位符残留: {text[:50]}...")
            
            # 操作指引残留
            if self.note_marker in text:
                results["note_residuals"] += 1
                if len(results["issues"]) < 5:
                    results["issues"].append(f"操作指引残留: {text[:50]}...")
            
            # Jinja2标签
            results["jinja2_tags"] += len(re.findall(r"\{\{.*?\}\}", text))
            results["jinja2_tags"] += len(re.findall(r"\{%.*?%\}", text))
            
            # 未闭合标签
            open_vars = len(re.findall(r"\{\{[^}]*$", text))
            open_blocks = len(re.findall(r"\{%[^%]*$", text))
            results["unclosed_tags"] += open_vars + open_blocks
        
        return results


# ============================================================
# 配置加载
# ============================================================

def load_config(config_path: Optional[str]) -> ReplacementConfig:
    """加载配置文件"""
    config = ReplacementConfig()
    
    if not config_path:
        return config
    
    if not os.path.exists(config_path):
        logger.warning(f"配置文件不存在: {config_path}，使用默认配置")
        return config
    
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            if config_path.endswith('.json'):
                data = json.load(f)
            elif config_path.endswith('.yaml') or config_path.endswith('.yml'):
                try:
                    import yaml
                    data = yaml.safe_load(f)
                except ImportError:
                    logger.warning("未安装PyYAML，无法加载YAML配置")
                    return config
            else:
                logger.warning(f"不支持的配置文件格式: {config_path}")
                return config
        
        # 更新配置
        if 'global_replacements' in data:
            config.global_replacements = [
                tuple(item) if isinstance(item, list) else item
                for item in data['global_replacements']
            ]
        if 'placeholder_pattern' in data:
            config.placeholder_pattern = data['placeholder_pattern']
        if 'placeholder_template' in data:
            config.placeholder_template = data['placeholder_template']
        if 'note_pattern' in data:
            config.note_pattern = data['note_pattern']
        if 'infer_descriptions' in data:
            config.infer_descriptions = data['infer_descriptions']
        if 'context_length' in data:
            config.context_length = data['context_length']
        
        logger.info(f"已加载配置文件: {config_path}")
        
    except Exception as e:
        logger.error(f"加载配置文件失败: {e}")
    
    return config


# ============================================================
# CLI入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Word模板转换器 - 将Word文档转换为docxtpl模板",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 基本使用
  python word_template_converter.py input.docx output.docx

  # 使用配置文件
  python word_template_converter.py --config config.json input.docx output.docx

  # 仅验证
  python word_template_converter.py --verify input.docx

配置文件格式 (config.json):
  {
    "global_replacements": [
      ["公司名称", "{{ g.company_name }}"],
      ["报告编号", "{{ g.report_no }}"]
    ],
    "infer_descriptions": true,
    "context_length": 60
  }
        """,
    )
    parser.add_argument("input", help="输入文件路径（原始Word文档）")
    parser.add_argument("output", nargs="?", help="输出文件路径（docxtpl模板）")
    parser.add_argument("--config", "-c", help="配置文件路径（JSON或YAML）")
    parser.add_argument("--verify", action="store_true", help="仅验证，不输出")
    parser.add_argument("--verbose", "-v", action="store_true", help="详细输出")
    parser.add_argument("--replace", "-r", nargs=2, metavar=("OLD", "NEW"),
                       action="append", help="添加全局替换规则（可多次使用）")
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # 检查输入文件
    if not os.path.exists(args.input):
        logger.error(f"输入文件不存在: {args.input}")
        sys.exit(1)
    
    # 加载配置
    config = load_config(args.config)
    
    # 添加命令行指定的全局替换
    if args.replace:
        for old, new in args.replace:
            config.global_replacements.append((old, new))
    
    # 加载文档
    logger.info(f"加载文档: {args.input}")
    doc = Document(args.input)
    
    logger.info(f"文档统计: {len(doc.paragraphs)} 段落, {len(doc.tables)} 表格, {len(doc.sections)} 节")
    
    if args.verify:
        # 仅验证模式
        logger.info("执行验证...")
        validator = TemplateValidator()
        results = validator.validate(doc)
        
        print("\n" + "=" * 60)
        print("验证结果:")
        print(f"  占位符残留: {results['placeholder_residuals']}")
        print(f"  操作指引残留: {results['note_residuals']}")
        print(f"  Jinja2标签数量: {results['jinja2_tags']}")
        print(f"  未闭合标签: {results['unclosed_tags']}")
        
        if results['issues']:
            print("\n  发现的问题:")
            for issue in results['issues']:
                print(f"    - {issue}")
        print("=" * 60)
        
        # 判断是否通过
        passed = (
            results["placeholder_residuals"] == 0
            and results["note_residuals"] == 0
            and results["unclosed_tags"] == 0
        )
        
        if passed:
            print("✓ 验证通过！")
        else:
            print("✗ 验证未通过，存在残留问题。")
        sys.exit(0 if passed else 1)
    
    # 需要输出文件
    if not args.output:
        logger.error("请指定输出文件路径，或使用 --verify 仅验证")
        sys.exit(1)
    
    # 创建替换器和处理器
    replacer = TextReplacer(config)
    processor = DocumentProcessor(replacer)
    
    # 执行处理
    logger.info("执行替换操作...")
    stats = processor.process(doc)
    
    # 保存输出
    logger.info(f"保存模板: {args.output}")
    output_dir = os.path.dirname(args.output)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(args.output)
    
    # 输出统计
    print("\n" + "=" * 60)
    print("替换统计:")
    print(f"  正文段落处理: {stats.paragraphs_processed} 处")
    print(f"  表格单元格处理: {stats.tables_processed} 处")
    print(f"  页眉页脚处理: {stats.headers_footers_processed} 处")
    print(f"  全局变量替换: {stats.globals_replaced} 次")
    print(f"  占位符替换: {stats.placeholders_replaced} 个")
    print(f"  操作指引清理: {stats.notes_removed} 处")
    print("=" * 60)
    
    # 验证输出
    logger.info("验证输出文件...")
    doc_out = Document(args.output)
    validator = TemplateValidator()
    results = validator.validate(doc_out)
    
    print(f"\n验证结果:")
    print(f"  占位符残留: {results['placeholder_residuals']}")
    print(f"  操作指引残留: {results['note_residuals']}")
    print(f"  Jinja2标签: {results['jinja2_tags']}")
    print("=" * 60)
    
    logger.info(f"模板转换完成: {args.output}")


if __name__ == "__main__":
    main()
