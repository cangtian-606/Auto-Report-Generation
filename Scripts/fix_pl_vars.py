"""Fix PL table variable names - use bracket notation for all PL keys"""
from docx import Document

doc = Document('templates/FDD模板.docx')
t13 = doc.tables[13]

pl_items = ['营业收入', '营业成本', '税金及附加', '销售费用', '财务费用', '营业利润', '利润总额', '所得税费用', '净利润']
pl_periods = ['2026年1至4月', '2025年度', '2024年度']

for ri in range(1, len(t13.rows)):
    item = pl_items[ri - 1] if ri - 1 < len(pl_items) else None
    if not item:
        continue
    for ci in range(1, len(t13.columns)):
        period = pl_periods[ci - 1] if ci - 1 < len(pl_periods) else ''
        var_name = f"{{{{ date.利润表['{item}_{period}'] | money }}}}"
        cell = t13.cell(ri, ci)
        for p in cell.paragraphs:
            p.clear()
            p.add_run(var_name)
        if ci == 1:
            print(f'R{ri} C1: {var_name}')

doc.save('templates/FDD模板.docx')
print('Saved')
