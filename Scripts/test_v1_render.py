"""Quick render test for FDD v1 template"""
import sys; sys.path.insert(0, '.')
from src.yaml_reader import YamlDataReader
from src.generator import DocumentGenerator

reader = YamlDataReader('data/FDDv1_数据字典.yaml')
ctx = reader.read_context()

# Fill sample data
ctx['date']['全局'] = {
    '公司名称': '重庆晟和泰新能源科技有限公司',
    '公司简称': '重庆晟和泰',
    '报告基准日': '2026年4月30日',
    '报告基准日_大写': '二〇二六年四月三十日',
    '报告文号': '大华咨字[2026]0001号',
    '项目名称': '重庆宏声印务',
}
ctx['date']['历史沿革'] = {
    '公司名称': '重庆晟和泰新能源科技有限公司',
    '统一社会信用代码': '91500227MA5U7XXXXX',
    '注册地址': '重庆市璧山区璧泉街道XX路XX号',
    '法定代表人': '张某某',
    '公司类型': '有限责任公司',
    '注册资本': 500,
    '经营范围': '新能源技术开发、太阳能发电',
    '成立日期': '2021-05-20',
    '营业期限': '2021-05-20 至 无固定期限',
}
ctx['date']['项目概况'] = {
    '上网模式': '自发自用，余量上网',
    '不含税总价': 2681250,
    '不含税总价单价': 3.75,
    '地址': '重庆市璧山区',
    '竣工容量': 0.715,
    '装机容量': 0.80,
}
ctx['form']['释义'] = [
    {'全称': '重庆晟和泰新能源科技有限公司', '简称': '公司'},
    {'全称': '大华会计师事务所(特殊普通合伙)上海分所', '简称': '本所'},
    {'全称': '成都交投投资有限公司', '简称': '成都交投'},
]
ctx['form']['股权结构'] = [
    {'股东名称': '张伟', '认缴出资额': 300, '认缴比例': 0.60, '实缴出资额': 300, '实缴比例': 0.60},
    {'股东名称': '李芳', '认缴出资额': 200, '认缴比例': 0.40, '实缴出资额': 200, '实缴比例': 0.40},
]
ctx['form']['股东信息'] = [
    {'股东名称': '张伟', '股东介绍': '执行董事兼总经理'},
    {'股东名称': '李芳', '股东介绍': '监事'},
]
ctx['form']['公司设立'] = []

gen = DocumentGenerator('templates/FDD初稿 v1.docx')
gen.render(ctx, 'output/FDDv1_测试输出.docx')
print('RENDER OK')

from docx import Document
doc = Document('output/FDDv1_测试输出.docx')
all_text = ''
for p in doc.paragraphs: all_text += p.text + '\n'
for t in doc.tables:
    for r in t.rows:
        for c in r.cells: all_text += c.text + '\n'

has_tags = '{{' in all_text
print('Unrendered tags:', has_tags)

checks = ['重庆晟和泰', '91500227', '张伟', '500', '0.80', '2681250']
for c in checks:
    print('  [%s]: %s' % (c, 'OK' if c in all_text else 'MISSING'))
