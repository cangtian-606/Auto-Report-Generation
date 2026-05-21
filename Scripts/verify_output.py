from docx import Document
doc = Document('output/FDD_测试输出.docx')

t2 = doc.tables[2]
print(f'T2 R1 C1 credit_code: {t2.cell(1, 1).text[:50]}')
print(f'T2 R5 C1 capital: {t2.cell(5, 1).text[:50]}')

t7 = doc.tables[7]
print(f'T7 R4 C3 total: {t7.cell(4, 3).text[:50]}')

t10 = doc.tables[10]
print(f'T10 R1 C1: {t10.cell(1, 1).text[:30]}')
print(f'T10 R1 C4: {t10.cell(1, 4).text[:30]}')

t12 = doc.tables[12]
print(f'T12 R2 C1 AR: {t12.cell(2, 1).text[:30]}')

t13 = doc.tables[13]
print(f'T13 R9 C3 NP: {t13.cell(9, 3).text[:50]}')

all_text = ''
for p in doc.paragraphs:
    all_text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            all_text += c.text
has_tags = '{{' in all_text
print(f'Unrendered tags: {has_tags}')
print(f'Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}')
print('VERDICT: RENDER OK' if not has_tags else 'VERDICT: HAS UNRENDERED TAGS')
