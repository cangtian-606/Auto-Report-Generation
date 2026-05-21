"""Fix FDD v1 template - rewrite paragraphs and cells reliably"""
from pathlib import Path
from docx import Document

PROJECT = Path(__file__).resolve().parent.parent
SRC = PROJECT / "templates" / "FDD初稿 v1.docx"
OUT = PROJECT / "templates" / "FDD初稿 v1.docx"

doc = Document(str(SRC))

# ===== Fix 1: Table[1] 释义 ====
t1 = doc.tables[1]
for ri in range(len(t1.rows)):
    for ci in range(len(t1.columns)):
        cell = t1.cell(ri, ci)
        full = ''.join(r.text for r in cell.paragraphs[0].runs) if cell.paragraphs[0].runs else ''
        if '{{ i.序号 }}' in full:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.text = ''
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].text = '{{ i.全称 }}'
            else:
                cell.paragraphs[0].add_run('{{ i.全称 }}')
        if '{{ i.职务 }}' in full:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.text = ''
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].text = '{{ i.简称 }}'
            else:
                cell.paragraphs[0].add_run('{{ i.简称 }}')
print("Fix 1: T[1] col names")

# ===== Fix 3+4: P[108] and P[110] ====
for i, p in enumerate(doc.paragraphs):
    full = ''.join(r.text for r in p.runs)
    if not full.strip():
        continue

    # P[108]: add space before }} 
    if '{{ date.全局.项目名称}}项目位于' in full:
        new = full.replace('{{ date.全局.项目名称}}', '{{ date.全局.项目名称 }}')
        new = new.replace('{{ date.项目概况.装机容量}}', '{{ date.项目概况.装机容量 }}')
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = new
        print("Fix 3: P[%d] 格式统一" % i)

    # P[110]: space removal + 竣工容量
    if '竣工安装容量为{{ date.项目概况.装机容量}}' in full:
        new = full.replace(
            '竣工安装容量为{{ date.项目概况.装机容量}}',
            '竣工安装容量为{{ date.项目概况.竣工容量 }}'
        )
        new = new.replace(
            '{{ date.项目概况. 不含税总价单价}}',
            '{{ date.项目概况.不含税总价单价 }}'
        )
        new = new.replace('{{ date.全局.项目名称}}', '{{ date.全局.项目名称 }}')
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = new
        print("Fix 4: P[%d] 竣工容量 + 去空格" % i)

doc.save(str(OUT))

# ===== VERIFY =====
doc2 = Document(str(OUT))

t1 = doc2.tables[1]
print("\nT[1]:")
for ri in range(len(t1.rows)):
    cells = [t1.cell(ri, ci).text[:50] for ci in range(len(t1.columns))]
    print("  R[%d]: %s" % (ri, cells))

for i, p in enumerate(doc2.paragraphs):
    text = ''.join(r.text for r in p.runs) if p.runs else ''
    if '项目位于' in text or '竣工安装容量' in text:
        print("P[%d]: %s" % (i, text[:200]))
