"""Fix Chinese key names: use bracket notation in template"""
from docx import Document
import re

doc = Document('templates/FDD_项目概况模板.docx')

t7 = doc.tables[7]
for ri in range(1, len(t7.rows)):
    for ci in range(len(t7.columns)):
        cell = t7.cell(ri, ci)
        for p in cell.paragraphs:
            for r in p.runs:
                m = re.search(r'form\.PC合同\.(\S+)', r.text)
                if m:
                    key = m.group(1)
                    old = 'PC合同.' + key
                    new = "PC合同['" + key + "']"
                    r.text = r.text.replace(old, new)

# Also fix any other problematic dot-access patterns in all cells
for t in doc.tables:
    for ri in range(len(t.rows)):
        for ci in range(len(t.columns)):
            for p in t.cell(ri, ci).paragraphs:
                for r in p.runs:
                    if 'PC合同.' in r.text and "['" not in r.text:
                        r.text = r.text.replace('PC合同.', "PC合同['").replace(' | ', "'] | ")

# Check
for ri in range(1, len(t7.rows)):
    print('T7 R%d C3: %s' % (ri, t7.cell(ri, 3).text[:80]))

doc.save('templates/FDD_项目概况模板.docx')
print('Saved')
