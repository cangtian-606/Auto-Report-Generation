"""Fix loop table structure: for / data / endfor must be separate rows"""
from docx import Document

doc = Document('templates/FDD_项目概况模板.docx')

def set_cell_text(table, row, col, text):
    cell = table.cell(row, col)
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    cell.paragraphs[0].clear()
    cell.paragraphs[0].add_run(text)


# ===== T1: 用语释义：添加 for/endfor 标签 =====
t1 = doc.tables[1]
# Clear row1 col0 and insert for tag
old_text = t1.cell(1, 0).text
for p in t1.cell(1, 0).paragraphs:
    for r in p.runs: r.text = ''
t1.cell(1, 0).paragraphs[0].add_run('{%tr for i in form.用语释义 %}')
# row2 and row3 data are already set as {{ i.全称 }} / {{ i.简称 }}
# Add endfor at last row
last_row_idx = len(t1.rows) - 1
# Replace last row col0 with endfor
for p in t1.cell(last_row_idx, 0).paragraphs:
    for r in p.runs: r.text = ''
t1.cell(last_row_idx, 0).paragraphs[0].add_run('{%tr endfor %}')
for ci in range(1, len(t1.columns)):
    set_cell_text(t1, last_row_idx, ci, '')
print('T[1] fixed')

# ===== T3: 股东出资 =====
t3 = doc.tables[3]
# Row0 = header (keep)
# Row1: clear all, put only {%tr for %} in col0
for ci in range(len(t3.columns)):
    set_cell_text(t3, 1, ci, '')
t3.cell(1, 0).paragraphs[0].add_run('{%tr for i in form.股东出资 %}')
# Add data row
data_row = t3.add_row()
set_cell_text(t3, 2, 0, '{{ i.股东名称 }}')
set_cell_text(t3, 2, 1, '{{ i.认缴出资额 | money }}')
set_cell_text(t3, 2, 2, '{{ i.认缴比例 | percent }}')
set_cell_text(t3, 2, 3, '{{ i.实缴出资额 | money }}')
if len(t3.columns) > 4:
    set_cell_text(t3, 2, 4, '{{ i.实缴比例 | percent }}')
# Add endfor row
endfor_row = t3.add_row()
set_cell_text(t3, 3, 0, '{%tr endfor %}')
for ci in range(1, len(t3.columns)):
    set_cell_text(t3, 3, ci, '')
print(f'T[3] fixed: {len(t3.rows)} rows')
for ri in range(len(t3.rows)):
    cells = [t3.cell(ri, ci).text[:40] for ci in range(len(t3.columns))]
    print(f'  R[{ri}]: {cells}')

# ===== T5: 股东出资_历史 =====
t5 = doc.tables[5]
for ci in range(len(t5.columns)):
    set_cell_text(t5, 1, ci, '')
t5.cell(1, 0).paragraphs[0].add_run('{%tr for i in form.股东出资_历史 %}')
# Add data row
data_row = t5.add_row()
set_cell_text(t5, 2, 0, '{{ i.股东名称 }}')
set_cell_text(t5, 2, 1, '{{ i.认缴出资额 | money }}')
set_cell_text(t5, 2, 2, '{{ i.认缴比例 | percent }}')
set_cell_text(t5, 2, 3, '{{ i.实缴出资额 | money }}')
if len(t5.columns) > 4:
    set_cell_text(t5, 2, 4, '{{ i.实缴比例 | percent }}')
# Add endfor row
endfor_row = t5.add_row()
set_cell_text(t5, 3, 0, '{%tr endfor %}')
for ci in range(1, len(t5.columns)):
    set_cell_text(t5, 3, ci, '')
print(f'T[5] fixed: {len(t5.rows)} rows')
for ri in range(len(t5.rows)):
    cells = [t5.cell(ri, ci).text[:40] for ci in range(len(t5.columns))]
    print(f'  R[{ri}]: {cells}')

doc.save('templates/FDD_项目概况模板.docx')
print('Saved')
