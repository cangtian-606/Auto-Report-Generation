#!/usr/bin/env python3
"""将 FDD初稿.docx 转换为 docxtpl Jinja2 模板"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt

PROJECT_DIR = Path(__file__).resolve().parent.parent
SRC = PROJECT_DIR / "templates" / "FDD初稿.docx"
OUT = PROJECT_DIR / "templates" / "FDD模板.docx"


def convert_template():
    doc = Document(str(SRC))
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10)

    _process_tables(doc)
    _process_paragraphs(doc)

    doc.save(str(OUT))
    print(f"模板已保存: {OUT}")


def _replace_runs(para, mapping):
    """替换段落中所有run的文本"""
    for run in para.runs:
        for old, new in mapping.items():
            if old in run.text:
                run.text = run.text.replace(old, new)


def _process_paragraphs(doc):
    for p in doc.paragraphs:
        text = p.text

        if '大华咨字' in text:
            _replace_runs(p, {'大华咨字[2026]XXXX号': '{{ 全局.报告文号 }}'})

        if '系由【】出资组建' in text:
            _replace_runs(p, {
                '【】': '{{ 历史沿革.出资方 }}',
                '【】年【】月【】日': '{{ 历史沿革.设立日期 }}',
                '【】核发的': '{{ 历史沿革.核发机关 }}',
                '统一社会信用代码为【】': '统一社会信用代码为{{ 基本情况.信用代码 }}'
            })

        if '【】项目位于【】' in text:
            _replace_runs(p, {
                '【】项目位于【】，': '{{ 项目概况.项目简称 }}项目位于{{ 项目概况.项目地址 }}，',
                '实际装机容量为【】': '实际装机容量为{{ 项目概况.装机容量 | num }}',
            })

        if '不含税工程总价为【】' in text:
            _replace_runs(p, {
                '【】项目竣工安装容量为【】，': '{{ 项目概况.项目简称 }}项目竣工安装容量为{{ 项目概况.竣工容量 | num }}，',
                '不含税工程总价为【】元': '不含税工程总价为{{ 项目概况.工程总价 | money }}元',
                '不含税建设单价为【】': '不含税建设单价为{{ 项目概况.建设单价 | money }}',
            })

        if '合同总价为' in text and '分项价格如下' in text:
            _replace_runs(p, {
                '【】项目合同总价为': '{{ 项目概况.项目简称 }}项目合同总价为',
                '2,728,160.00': '{{ PC合同.合同总价 | money }}',
            })

        if '乙方（承包单位）：【】' in text:
            _replace_runs(p, {'【】': '{{ PC合同.承包单位 }}'})

        if '工程内容及规模：【】' in text:
            _replace_runs(p, {'【】': '{{ PC合同.工程内容 }}'})

        if '电费结算单价的【】%' in text or '分时电度用电价格的【】%' in text:
            _replace_runs(p, {'【】': '{{ EMC合同.结算折扣率 }}'})

        if text.startswith('②用能方：') and '【】' in text:
            _replace_runs(p, {'【】': '{{ EMC合同.用能方 }}'})

        if '自发自用电量占总发电量比例为' in text:
            _replace_runs(p, {'*': '{{ 运营情况.自发自用占比 | percent }}'})

        if '上网电费含税单价为' in text:
            _replace_runs(p, {
                '*元/kWh': '{{ 运营情况.上网电价 | money }}元/kWh',
                '*65%': '{{ 运营情况.结算折扣率 }}%'
            })

        if '20' in text and 'x年' in text and '电费未回款' in text:
            text2 = text
            text2 = text2.replace('202x年x月至x月', '{{ 运营情况.自发自用未回款期间 }}')
            text2 = text2.replace('202x年x月', '{{ 运营情况.余电上网未回款期间 }}')
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = text2

        if '二〇二六年五月【】日' in text:
            _replace_runs(p, {'【】': '{{ 全局.报告日期_日 }}'})


def _process_tables(doc):
    # ===== Table[0]: Cover page =====
    t0 = doc.tables[0]
    _cell_replace(t0, 0, 0, '【】', '{{ 全局.项目公司简称 }}')
    _cell_replace(t0, 2, 0, '大华咨字[2026]XXXX号', '{{ 全局.报告文号 }}')

    # ===== Table[1]: Company full/short name mapping =====
    t1 = doc.tables[1]
    for ri in range(1, len(t1.rows)):
        _cell_replace_all(t1, ri, 0, 'xxxx', '{{ 全局.项目公司全称 }}')
        _cell_replace_all(t1, ri, 1, 'xxxx', '{{ 全局.项目公司简称 }}')

    # ===== Table[2]: Company basic info =====
    t2 = doc.tables[2]
    field_map_t2 = {
        '公司名称': '{{ 基本情况.公司名称 }}',
        '统一社会信用代码': '{{ 基本情况.信用代码 }}',
        '注册地址': '{{ 基本情况.注册地址 }}',
        '法定代表人': '{{ 基本情况.法定代表人 }}',
        '公司类型': '{{ 基本情况.公司类型 }}',
        '注册资本': '{{ 基本情况.注册资本 | num }}',
        '经营范围': '{{ 基本情况.经营范围 }}',
        '成立日期': '{{ 基本情况.成立日期 }}',
        '营业期限': '{{ 基本情况.营业期限 }}',
    }
    for ri in range(len(t2.rows)):
        key = t2.cell(ri, 0).text.strip()
        if '公司名称' in key and ri < len(t2.rows):
            # Replace value cell xxxx
            _cell_replace_all(t2, ri, 1, 'xxxx', field_map_t2.get(key, ''))

    # More generalized: for all rows replace xxxx in col 1
    for ri in range(len(t2.rows)):
        cell_text = t2.cell(ri, 1).text.strip()
        if 'xxxx' in cell_text:
            key = t2.cell(ri, 0).text.strip()
            var = field_map_t2.get(key, '{{ 基本情况.' + key + ' }}')
            _cell_replace_all(t2, ri, 1, 'xxxx', var)

    # ===== Table[3]: Shareholder table → loop =====
    t3 = doc.tables[3]
    _make_shareholder_loop(t3, '股东出资')

    # ===== Table[4]: Duplicate company info (another entity?) =====
    # Skip if it's a duplicate of Table[2] pattern

    # ===== Table[5]: Duplicate shareholder table =====
    t5 = doc.tables[5]
    _make_shareholder_loop(t5, '股东出资_历史')

    # ===== Table[6]: Project approvals =====
    t6 = doc.tables[6]
    for ri in range(len(t6.rows)):
        _cell_replace_all(t6, ri, 0, '【】', '{{ 全局.项目公司简称 }}')

    # ===== Table[7]: PC contract cost =====
    t7 = doc.tables[7]
    _cell_replace_all(t7, 1, 0, 'xxxx', '1')
    _cell_replace_all(t7, 1, 1, 'xxxx', '设备费')
    _cell_replace_all(t7, 1, 2, 'xxxx', '13%')
    _cell_replace_all(t7, 1, 3, 'xxxx', '{{ PC合同.设备费 | money }}')
    _cell_replace_all(t7, 2, 0, 'xxxx', '2')
    _cell_replace_all(t7, 2, 1, 'xxxx', '建筑安装工程费')
    _cell_replace_all(t7, 2, 2, 'xxxx', '9%')
    _cell_replace_all(t7, 2, 3, 'xxxx', '{{ PC合同.建安费 | money }}')
    _cell_replace_all(t7, 3, 0, 'xxxx', '3')
    _cell_replace_all(t7, 3, 1, 'xxxx', '其他费')
    _cell_replace_all(t7, 3, 2, 'xxxx', '6%')
    _cell_replace_all(t7, 3, 3, 'xxxx', '{{ PC合同.其他费 | money }}')
    _cell_replace_all(t7, 4, 0, 'xxxx', '合计')
    _cell_replace_all(t7, 4, 3, 'xxxx', '{{ PC合同.合同总价 | money }}')

    # ===== Table[8]: Tax rates → keep as-is =====
    # Fixed tax rates, no replacement needed

    # ===== Table[9]: Fixed assets params =====
    t9 = doc.tables[9]
    _cell_replace_all(t9, 0, 2, '【】项目', '{{ 全局.项目公司简称 }}项目')
    field_vals = {
        '固定资产原值': '{{ 固定资产.原值 | money }}',
        '残值率': '{{ 固定资产.残值率 | percent }}',
        '折旧年限': '{{ 固定资产.折旧年限 }}',
        '转固时点': '{{ 固定资产.转固时点 }}',
        '年折旧额': '{{ 固定资产.年折旧额 | money }}',
    }
    for ri in range(1, len(t9.rows)):
        key = t9.cell(ri, 0).text.replace(str(ri), '').strip()
        if key in field_vals:
            _cell_replace_all(t9, ri, 2, '[]', field_vals[key])

    # ===== Table[10]: Power generation → keys per period =====
    t10 = doc.tables[10]
    periods = ['2024年度', '2025年度', '2026年1-4月']
    for ri in range(1, len(t10.rows)):
        for ci in range(len(t10.columns)):
            _cell_replace_all(t10, ri, ci, '[]', '')

    # ===== Table[11]: O&M contract =====
    t11 = doc.tables[11]
    field_vals_t11 = {
        '服务期限': '{{ 运维合同.服务期限 }}',
        '合同金额（含税）': '{{ 运维合同.合同金额 | money }}',
        '合同单价（含税）': '{{ 运维合同.合同单价 | money }}',
        '进项税率': '{{ 运维合同.进项税率 }}',
    }
    for ri in range(len(t11.rows)):
        key = t11.cell(ri, 0).text.strip()
        if key in field_vals_t11:
            t11.cell(ri, 1).text = field_vals_t11[key]

    # ===== Table[12]: Balance sheet =====
    t12 = doc.tables[12]
    _process_financial_table(t12, 'date.资产负债表')

    # ===== Table[13]: Income statement =====
    t13 = doc.tables[13]
    _process_financial_table(t13, 'date.利润表')

    # ===== Table[14]: Related parties → loop =====
    t14 = doc.tables[14]
    if len(t14.rows) > 1:
        _make_cell_for_loop(t14, 0, '关联方')

    # ===== Table[15]: Related party balances → loop =====
    t15 = doc.tables[15]
    if len(t15.rows) > 1:
        _make_cell_for_loop(t15, 0, '关联方余额')

    # ===== Table[16]: Related party transactions → loop =====
    t16 = doc.tables[16]
    if len(t16.rows) > 1:
        _make_cell_for_loop(t16, 0, '关联方交易')


def _cell_replace(table, row, col, old, new):
    for para in table.cell(row, col).paragraphs:
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)


def _cell_replace_all(table, row, col, old, new):
    for para in table.cell(row, col).paragraphs:
        text = ''.join(r.text for r in para.runs)
        if old in text:
            for r in para.runs:
                r.text = ''
            if para.runs:
                para.runs[0].text = text.replace(old, new)


def _make_shareholder_loop(table, loop_name):
    if len(table.rows) <= 1:
        return
    # Row 0 = header, keep as is
    # Add for row, data row, endfor row
    _clear_cell_keep_first_run(table, 1, 0)
    table.cell(1, 0).paragraphs[0].runs[0].text = '{%tr for i in ' + loop_name + ' %}' if table.cell(1, 0).paragraphs[0].runs else '{%tr for i in ' + loop_name + ' %}'
    for ci in range(1, table.rows[1].cells.__len__()):
        _clear_cell(table, 1, ci)

    if len(table.rows) > 2:
        _replace_all_runs(table.cell(2, 0), 'xxxx', '{{ i.股东名称 }}')
        _replace_all_runs(table.cell(2, 1), 'xxxx', '{{ i.认缴出资额 | money }}')
        _replace_all_runs(table.cell(2, 2), 'xxxx', '{{ i.认缴比例 | percent }}')
        _replace_all_runs(table.cell(2, 3), 'xxxx', '{{ i.实缴出资额 | money }}')
        if len(table.columns) > 4:
            _replace_all_runs(table.cell(2, 4), 'xxxx', '{{ i.实缴比例 | percent }}')

    # Add endfor row
    if len(table.rows) > 3:
        _clear_cell_keep_first_run(table, 3, 0)
        table.cell(3, 0).paragraphs[0].runs[0].text = '{%tr endfor %}' if table.cell(3, 0).paragraphs[0].runs else '{%tr endfor %}'
        for ci in range(1, table.rows[3].cells.__len__()):
            _clear_cell(table, 3, ci)


def _make_cell_for_loop(table, col_index, loop_name):
    if len(table.rows) <= 1:
        return
    _clear_cell(table, 1, col_index)
    p = table.cell(1, col_index).paragraphs[0]
    p.add_run('{%tr for i in ' + loop_name + ' %}')

    # Add endfor
    end_row_idx = min(3, len(table.rows) - 1)
    _clear_cell(table, end_row_idx, col_index)
    p = table.cell(end_row_idx, col_index).paragraphs[0]
    p.add_run('{%tr endfor %}')


def _clear_cell(table, row, col):
    for p in table.cell(row, col).paragraphs:
        for r in p.runs:
            r.text = ''


def _clear_cell_keep_first_run(table, row, col):
    cell = table.cell(row, col)
    for pi, p in enumerate(cell.paragraphs):
        if pi == 0:
            for r in p.runs:
                r.text = ''
            break
        for r in p.runs:
            r.text = ''


def _replace_all_runs(cell, old, new):
    for p in cell.paragraphs:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)


def _process_financial_table(table, prefix):
    for ri in range(1, len(table.rows)):
        field = table.cell(ri, 0).text.strip()
        if not field:
            continue
        var_name = prefix + '.' + field
        for ci in range(1, len(table.columns)):
            _cell_replace_all(table, ri, ci, '', '')


if __name__ == '__main__':
    convert_template()
