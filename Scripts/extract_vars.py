"""Final comprehensive fix + extract all variables for data dictionary"""
from pathlib import Path
from docx import Document
import re

PROJECT = Path(__file__).resolve().parent.parent
SRC = PROJECT / "templates" / "FDD初稿 v1.docx"

doc = Document(str(SRC))

# ===== Fix all spacing issues: \.xxx}} → .xxx }} =====
# Remove space after dot in variable names
# Add space before }}
for i, p in enumerate(doc.paragraphs):
    full = ''.join(r.text for r in p.runs)
    if '{{' not in full:
        continue
    
    changed = False
    new_text = full
    
    # Fix: . xxx}} → .xxx}} (remove space after dot)
    new_text = re.sub(r'\.\s+(\w+)', r'.\1', new_text)
    
    # Fix: xxx}} → xxx }} (add space before closing brackets - only when not already)
    new_text = re.sub(r'(\w)\}\}', r'\1 }}', new_text)
    # But don't double-space: }} → }} (catch duplicate)
    new_text = new_text.replace(' }} }}', ' }}')
    
    if new_text != full:
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = new_text
        changed = True

# Same for tables
for t in doc.tables:
    for ri in range(len(t.rows)):
        for ci in range(len(t.columns)):
            cell = t.cell(ri, ci)
            full = ''.join(r.text for p in cell.paragraphs for r in p.runs)
            if '{{' not in full:
                continue
            new_text = full
            new_text = re.sub(r'\.\s+(\w+)', r'.\1', new_text)
            new_text = re.sub(r'(\w)\}\}', r'\1 }}', new_text)
            new_text = new_text.replace(' }} }}', ' }}')
            if new_text != full:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = ''
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].text = new_text
                else:
                    cell.paragraphs[0].add_run(new_text)

# Fix headers
for section in doc.sections:
    for header in [section.header]:
        if not header: continue
        for p in header.paragraphs:
            for r in p.runs:
                r.text = r.text.replace('{{ date.全局.报告文号 }}号', '{{ date.全局.报告文号 }}号')  # OK, keep

doc.save(str(SRC))
print("Format unified")

# ===== EXTRACT ALL VARIABLES =====
doc2 = Document(str(SRC))
all_text = ''
for p in doc2.paragraphs:
    all_text += ''.join(r.text for r in p.runs) + '\n'
for t in doc2.tables:
    for r in t.rows:
        for c in r.cells:
            all_text += c.text + '\n'
for s in doc2.sections:
    h = s.header
    if h:
        all_text += ' '.join(p.text for p in h.paragraphs) + '\n'

# Find all {{ date.xxx.yyy }} patterns
date_vars = set()
form_vars = set()
for m in re.finditer(r'\{\{\s*date\.([a-zA-Z_0-9\u4e00-\u9fff]+(\.([a-zA-Z_0-9\u4e00-\u9fff]+))*)', all_text):
    date_vars.add(m.group(1))
for m in re.finditer(r'\{\%\s*for\s+i\s+in\s+form\.([a-zA-Z_0-9\u4e00-\u9fff]+)\s*\%\}', all_text):
    form_vars.add(m.group(1))

print("\n=== date variables ===")
for v in sorted(date_vars):
    print("  %s" % v)

print("\n=== form tables ===")
for v in sorted(form_vars):
    print("  %s" % v)
    # Find {{ i.xxx }} references near this table
    i_vars = set()
    for m in re.finditer(r'\{\{\s*i\.([a-zA-Z_0-9\u4e00-\u9fff]+)', all_text):
        i_vars.add(m.group(1))
    print("    columns: %s" % sorted(i_vars))

print("\nTotal date vars:", len(date_vars))
print("Total form tables:", len(form_vars))
