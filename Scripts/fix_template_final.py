"""
修复尽调报告模板:
1. 消除 for 块级循环空行: 把 for/endfor 标签合并到前后段落
2. 手续 tc 表每行加 tc 列循环
3. 条件控制列: 通过预过滤项目列表
"""
import shutil, zipfile
from pathlib import Path
from docx import Document

project = Path(r'D:\ProgramWorkingSpace\Auto-Report-Generation')
src = project / 'templates' / '尽调报告模板v1.docx'
backup = project / 'templates' / '尽调报告模板v1_backup.docx'

if not backup.exists():
    shutil.copy2(str(src), str(backup))
    print("Backup saved: " + backup.name)

doc = Document(str(src))

# ============================================================
# Fix 1: 消除 {% for %} 块级循环空行
# 策略: 把 for 标签移到前一段落末尾, endfor 移到后一段落开头
# 这样 for/endfor 不与空内容共享段落, 不会产生空行
# ============================================================

# 找到目标段落索引
for_first_i = None
endfor_first_i = None
content_first_i = None
for_second_i = None
endfor_second_i = None
content_second_i = None

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == '{% for i in 项目基本情况.项目公司 %}':
        if for_first_i is None:
            for_first_i = i
        else:
            for_second_i = i
    elif t == '{% endfor %}':
        if endfor_first_i is None:
            endfor_first_i = i
        else:
            endfor_second_i = i

if for_first_i is not None:
    content_first_i = for_first_i + 1
if for_second_i is not None:
    content_second_i = for_second_i + 1

print("Block 1: for=P{} content=P{} endfor=P{}".format(for_first_i, content_first_i, endfor_first_i))
print("Block 2: for=P{} content=P{} endfor=P{}".format(for_second_i, content_second_i, endfor_second_i))

# Fix: append "{% for i in 项目基本情况.项目公司 %}" to the paragraph BEFORE for,
# and prepend "{% endfor %}" to the paragraph AFTER endfor.
# Then clear the for and endfor paragraphs.

# Block 1: P103 "项目概况" ← append for tag, P107 "项目总投资" ← prepend endfor
# Block 2: P107 already handled, P111 "项目相关手续办理情况" ← prepend endfor

# Get paragraph objects
for_p1 = doc.paragraphs[for_first_i]
endfor_p1 = doc.paragraphs[endfor_first_i]
content_p1 = doc.paragraphs[content_first_i]
prev_p1 = doc.paragraphs[for_first_i - 1] if for_first_i > 0 else None  # P103 "项目概况"
next_p1 = doc.paragraphs[endfor_first_i + 1] if endfor_first_i + 1 < len(doc.paragraphs) else None  # P107 "项目总投资"

if prev_p1 is not None:
    # Append for tag to end of P103
    old_text = prev_p1.text
    prev_p1.clear()
    prev_p1.add_run(old_text + '{% for i in 项目基本情况.项目公司 %}')
    print("  Block1: appended for tag to P{}".format(for_first_i - 1))

# Clear the for paragraph (P104)
for_p1.clear()

# Keep P105 (content) as-is - it's the content paragraph inside the for block

if next_p1 is not None:
    # Prepend endfor tag to start of P107
    old_text = next_p1.text
    next_p1.clear()
    next_p1.add_run('{% endfor %}' + old_text)
    print("  Block1: prepended endfor tag to P{}".format(endfor_first_i + 1))

# Clear the endfor paragraph (P106)
endfor_p1.clear()

# Block 2: Same pattern
if for_second_i is not None:
    # P108 is for, P109 is content, P110 is endfor
    for_p2 = doc.paragraphs[for_second_i]
    endfor_p2 = doc.paragraphs[endfor_second_i]
    prev_p2 = doc.paragraphs[for_second_i - 1] if for_second_i > 0 else None  # P107 "项目总投资"
    next_p2 = doc.paragraphs[endfor_second_i + 1] if endfor_second_i + 1 < len(doc.paragraphs) else None  # P111

    if prev_p2 is not None:
        old_text = prev_p2.text
        prev_p2.clear()
        prev_p2.add_run(old_text + '{% for i in 项目基本情况.项目公司 %}')
        print("  Block2: appended for tag to P{}".format(for_second_i - 1))

    for_p2.clear()

    if next_p2 is not None:
        old_text = next_p2.text
        next_p2.clear()
        next_p2.add_run('{% endfor %}' + old_text)
        print("  Block2: prepended endfor tag to P{}".format(endfor_second_i + 1))

    endfor_p2.clear()

# ============================================================
# Fix 2: 手续 tc 表 — 每行加 tc 列循环 + 条件控制
# ============================================================
# R00: 序号 | 阶段划分 | {%tc for i in 项目基本情况.项目公司 %} | {{ i.项目名称 }} | {%tc endfor %}
# R01-R12: 需要对每行也加 tc for, 且能按项目控制列显隐
#
# 简化方案: 
#   - R00 不变
#   - R01+ 每行改成: {%tc for i in 项目基本情况.项目公司 %}
#   - 每个项目加 '显示在手续表' 字段控制, 哪列需要显示

# 找 Table 5 (索引5)
tbl = doc.tables[5]
print("\n=== Fixing Table 5 (手续) ===")
print("Rows: {}, Cols: {}".format(len(tbl.rows), len(tbl.columns)))

# R01-R12: 把固定内容 "是" 改为 tc 列循环
# 当前 R01: 一 | 前期审批文件 | [empty] | [empty] | [empty] | [empty]
# 改成:   一 | 前期审批文件 | {%tc for i in 项目基本情况.项目公司 %} | {{ item[i.项目名称] }} | {%tc endfor %}
# 但这样需要数据结构支持: item = {"项目1": "是", "项目2": "否"}

# 更简单的方案: 使用键值对
# R01 tc: {{ item['手续_前期审批文件'][i.项目名称] }}
# 在 context 中: "手续_前期审批文件": {"项目1": "是", "项目2": "是"}

# 对于条件控制列显隐:
# 方案1: 在 context 中放 "项目公司_手续" 列表, 只包含需要显示的项目
# 方案2: 每个项目加 '显示在手续表' bool 字段

# 我们先实现方案1 (预过滤列表)

# 把 R00 改成使用预过滤列表
r0 = tbl.rows[0]
r0_c3 = r0.cells[2].paragraphs[0]
# 替换: {%tc for i in 项目基本情况.项目公司%} → {%tc for i in 项目基本情况.项目公司_手续%}
for run in r0_c3.runs:
    if '项目公司' in run.text:
        run.text = run.text.replace('项目公司', '项目公司_手续')
        print("  R00: using filtered list '项目公司_手续'")

# R01-R12: 加入 tc 循环
for ri in range(1, len(tbl.rows)):
    row = tbl.rows[ri]
    # Cell 2 = tc for start
    cell2 = row.cells[2]
    # Cell 3 = item value
    cell3 = row.cells[3]
    # Cell 4 = tc endfor
    cell4 = row.cells[4]

    # Clear and set
    cell2.text = ''
    cell3.text = ''
    cell4.text = ''

    # Set tc for start
    cell2.paragraphs[0].text = '{%tc for p in 项目基本情况.项目公司_手续%}'
    # Set value (use a naming convention)
    cell3.paragraphs[0].text = '{{ 手续.R{:02d}[p.项目名称] }}'.format(ri)
    # Set tc endfor
    cell4.paragraphs[0].text = '{%tc endfor%}'

print("  R01-R{}: added tc loops".format(len(tbl.rows) - 1))

# ============================================================
# Save
# ============================================================
doc.save(str(src))
print("\nTemplate saved: " + str(src))
print("Backup: " + str(backup))
print("\nDone! Now regenerate data with filtered project list...")
