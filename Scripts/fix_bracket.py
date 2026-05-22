"""Fix bracket notation for power generation table variables"""
from docx import Document

doc = Document('templates/FDD模板.docx')

t10 = doc.tables[10]
period_names = ['2024年度', '2025年度', '2026年1至4月']
col_types = ['自发自用', '余电上网', '发电量小计', '自发自用占比']

for ri in range(1, min(len(t10.rows), len(period_names) + 1)):
    pn = period_names[ri - 1]
    for ci in range(1, min(len(t10.columns), len(col_types) + 1)):
        ct = col_types[ci - 1]
        ft = 'num' if ci < 4 else 'percent'
        cell = t10.cell(ri, ci)
        for p in cell.paragraphs:
            p.clear()
            p.add_run(f"{{{{ 发电量['{pn}']['{ct}'] | {ft} }}}}")

# verify
for ri in range(1, min(len(t10.rows), 4)):
    for ci in range(1, 5):
        print(f'T10 R{ri} C{ci}: {t10.cell(ri, ci).text[:80]}')

doc.save('templates/FDD模板.docx')
print('Saved')
