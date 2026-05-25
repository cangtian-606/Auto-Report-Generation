"""删除模板中已被清空但仍残留的 for/endfor 段落 (P104, P106, P108, P110)
这些空段落仍在 {% for %} 块体内, 会在每轮迭代中产生空行"""
import zipfile, io, re, shutil
from pathlib import Path

project = Path(r'D:\ProgramWorkingSpace\Auto-Report-Generation')
src = project / 'templates' / '尽调报告模板v1.docx'

# Read XML
with zipfile.ZipFile(str(src), 'r') as zin:
    doc_xml = zin.read('word/document.xml').decode('utf-8')
    all_files = {fn: zin.read(fn) for fn in zin.namelist()}

# Find empty paragraphs between for blocks
# P104, P106 (project overview), P108, P110 (project investment)
# These are w:p elements with no text content (all cleared to EMPTY)
# We need to identify them by content and position

# Strategy: find </w:p> elements that contain NO significant text
# Simpler: find paragraphs that were cleared by searching for 
# specific pattern: empty paragraph between for content and endfor

# Actually, we should check the output to see if empty lines are still a problem.
# The test showed "[1]" and "[2]" - but did they have empty lines between them?
# Let me check the actual paragraph indices in output

from docx import Document as DocxDocument
doc = DocxDocument(str(src))

# Find all paragraphs for region P103-P111
print("Template fix verification:")
for i in [103, 104, 105, 106, 107, 108, 109, 110, 111]:
    t = doc.paragraphs[i].text
    status = "EMPTY" if not t.strip() else "HAS_TEXT"
    if status == "HAS_TEXT":
        print("  P{:03d}: {} ...{}".format(i, t[:70], t[-30:] if len(t) > 70 else ""))
    else:
        print("  P{:03d}: [{} - will be empty in output]".format(i, status))

# The problem: P104[EMPTY] and P106[EMPTY] ARE between {% for %} (in P103) and {% endfor %} (in P107)
# So {% for %} iterates over: P104(empty), P105(content), P106(empty)
# Each iteration produces 2 empty paragraphs + 1 content paragraph = 3 paragraphs
# We need to REMOVE P104 and P106 from the document

# XML approach: remove the w:p elements corresponding to P104, P106, P108, P110
from lxml import etree

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

tree = etree.fromstring(doc_xml.encode('utf-8'))
body = tree.find('.//{' + nsmap['w'] + '}body')
if body is None:
    # Try without namespace
    for child in tree:
        if child.tag.endswith('body'):
            body = child
            break

if body is None:
    print("ERROR: body not found")
else:
    w_p_tag = '{' + nsmap['w'] + '}p'
    paragraphs = list(body.findall(w_p_tag))
    print("\nTotal paragraphs: {}".format(len(paragraphs)))
    
    # Identify P104, P106, P108, P110 by their index + content
    # They should be empty paragraphs adjacent to content paragraphs
    to_remove = []
    
    for i, p in enumerate(paragraphs):
        text = ''.join(p.itertext()).strip()
        if not text:
            # Empty paragraph - check if it's between for content
            if i >= 2 and i + 1 < len(paragraphs):
                prev_text = ''.join(paragraphs[i-1].itertext()).strip()
                next_text = ''.join(paragraphs[i+1].itertext()).strip()
                
                # P104: prev=P103('项目概况{%for%}'), next=P105(content)  
                # P106: prev=P105(content), next=P107('{%endfor%}项目总投资%s%}')
                if '项目基本情况' in prev_text and 'for' in prev_text and '项目名称' in next_text:
                    to_remove.append(p)
                    print("  Will remove P104-like (between for-open and content): {}".format(i))
                elif '项目名称' in prev_text and 'endfor' in next_text and '项目总投资' in next_text:
                    to_remove.append(p)
                    print("  Will remove P106-like (between content and endfor+title): {}".format(i))
                elif prev_text and 'endfor' in prev_text and 'for' in prev_text and '含税总价' in next_text:
                    to_remove.append(p)
                    print("  Will remove P108-like (between for-open and content2): {}".format(i))
                elif '含税总价' in prev_text and 'endfor' in next_text and '手续' in next_text:
                    to_remove.append(p)
                    print("  Will remove P110-like (between content2 and endfor): {}".format(i))
    
    for p in to_remove:
        body.remove(p)
    
    print("Removed {} empty paragraphs".format(len(to_remove)))
    
    # Save
    new_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>' + etree.tostring(tree, encoding='unicode')
    
    with zipfile.ZipFile(str(src), 'w', zipfile.ZIP_DEFLATED) as zout:
        for fn, data in all_files.items():
            if fn == 'word/document.xml':
                zout.writestr(fn, new_xml.encode('utf-8'))
            else:
                zout.writestr(fn, data)
    
    print("Template saved")

# Final verification
print("\nFinal check:")
doc2 = DocxDocument(str(src))
for i, p in enumerate(doc2.paragraphs):
    t = p.text
    if '项目概况' in t or ('for' in t and 'endfor' not in t) or '项目名称' in t or 'endfor' in t or '项目总投资' in t or '含税总价' in t or '手续' in t:
        if t.strip():
            print("  P{:03d}: {}".format(i, t[:100]))
        else:
            print("  P{:03d}: [EMPTY]".format(i))
