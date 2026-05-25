"""分析 for 循环周边 + tc 表完整结构"""
from docx import Document

path = r'D:\ProgramWorkingSpace\Auto-Report-Generation\templates\尽调报告模板v1.docx'
doc = Document(path)

print('=== 段落 99-116 ===')
for i in range(99, min(117, len(doc.paragraphs))):
    t = doc.paragraphs[i].text
    if not t.strip():
        print("P{:03d}: [EMPTY]".format(i))
    else:
        print("P{:03d}: {}".format(i, t[:130]))

print()
for ti, tbl in enumerate(doc.tables):
    if ti == 5:
        print("Table 5: {} rows x {} cols".format(len(tbl.rows), len(tbl.columns)))
        for ri, row in enumerate(tbl.rows):
            cells = [c.text.strip()[:40] for c in row.cells]
            combined = " | ".join(cells)
            has_tag = "{%" in combined or "{{" in combined
            if has_tag or ri < 4:
                tag = " [TAG]" if has_tag else ""
                print("R{:02d}{}: {}".format(ri, tag, combined[:200]))
