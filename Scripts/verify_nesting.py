from docx import Document
doc = Document('output/嵌套测试输出.docx')

all_text = ''
for p in doc.paragraphs:
    all_text += p.text + '\n'
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            all_text += c.text + ' | '
        all_text += '\n'

print('=== FULL OUTPUT ===')
print(all_text[:3000])
print(f'\nHas unrendered tags: {"{{" in all_text}')
print(f'Contains 重庆晟和泰: {"重庆晟和泰" in all_text}')
print(f'Contains 安徽富军: {"安徽富军" in all_text}')
print(f'Contains 张伟: {"张伟" in all_text}')
print(f'Contains 赵丽: {"赵丽" in all_text}')
print(f'Contains 报告结束: {"报告结束" in all_text}')
