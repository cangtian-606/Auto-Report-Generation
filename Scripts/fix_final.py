"""Final fix for remaining xxxx/【】 placeholders in FDD template"""
from docx import Document

doc = Document('templates/FDD_项目概况模板.docx')

def set_cell_text(table, row, col, text):
    cell = table.cell(row, col)
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    cell.paragraphs[0].clear()
    cell.paragraphs[0].add_run(text)

# T4: clear xxxx (unidentified second entity table)
t4 = doc.tables[4]
for ri in range(len(t4.rows)):
    for ci in range(len(t4.columns)):
        ct = t4.cell(ri, ci).text
        if 'xxxx' in ct:
            set_cell_text(t4, ri, ci, '')

# T9 header column: 【】项目
t9 = doc.tables[9]
for p in t9.cell(0, 2).paragraphs:
    full = ''.join(r.text for r in p.runs)
    if '【】' in full:
        for r in p.runs: r.text = ''
        if p.runs:
            p.runs[0].text = full.replace('【】', '{{ date.全局.公司简称 }}')

doc.save('templates/FDD_项目概况模板.docx')

# Final check
doc2 = Document('templates/FDD_项目概况模板.docx')
all_text = ''
for p in doc2.paragraphs:
    all_text += ''.join(r.text for r in p.runs) + '\n'
for t in doc2.tables:
    for r in t.rows:
        for c in r.cells:
            all_text += c.text + '\n'

remaining = [c for c in ['【】', 'xxxx'] if c in all_text]
print('Remaining placeholders:', remaining)
jinja_count = all_text.count('{{')
print('Jinja2 var count:', jinja_count)
