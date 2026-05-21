"""Add outer {% for company in form.项目公司 %} loop to FDD template"""
from docx import Document
from docx.oxml.ns import qn

doc = Document('templates/FDD模板.docx')

# Find where per-company section starts: after "投资项目概况" heading
heading_indices = []
for i, p in enumerate(doc.paragraphs):
    if '投资项目概况' in p.text and 'Heading' in (p.style.name or ''):
        heading_indices.append(i)

# Find where conclusion starts: "重要的表外信息" or before signature
end_indices = []
for i, p in enumerate(doc.paragraphs):
    if '重要的表外信息' in p.text and 'Heading' in (p.style.name or ''):
        end_indices.append(i)
    if '大华会计师事务所' in p.text:
        end_indices.append(i)

if not heading_indices or not end_indices:
    # Fallback: find the start of company-specific content
    # After preamble, at "投资项目概况" section
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith('投资项目概况'):
            heading_indices.append(i)
            break

if not end_indices:
    # Use signature as end marker
    for i, p in enumerate(doc.paragraphs):
        if '大华会计师事务所' in p.text:
            end_indices.append(i)
            break

if heading_indices and end_indices:
    start_idx = heading_indices[0]
    end_idx = end_indices[0]

    # Access document body XML
    body = doc.element.body
    children = list(body)

    # Find the paragraph and table elements around start/end
    para_elements = []
    for i, p in enumerate(doc.paragraphs):
        para_elements.append(p._p)

    # Insert {% for %} before the heading paragraph
    for_para = doc.add_paragraph()
    for_para.style = doc.styles['Normal']
    for_run = for_para.add_run('{% for company in form.项目公司 %}')
    for_run.font.name = '宋体'

    # Insert {% endfor %} before the signature paragraph
    endfor_para = doc.add_paragraph()
    endfor_para.style = doc.styles['Normal']
    endfor_run = endfor_para.add_run('{%tr endfor %}')
    endfor_run.font.name = '宋体'

    # Move elements in body
    start_p = para_elements[start_idx] if start_idx < len(para_elements) else None
    end_p = para_elements[end_idx] if end_idx < len(para_elements) else None

    if start_p is not None and end_p is not None:
        body.insert(list(body).index(start_p), for_para._p)
        # Find end_p position after insertion
        end_pos = list(body).index(end_p)
        body.insert(end_pos, endfor_para._p)

    # Replace shareholder loop references inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if 'form.股东出资' in r.text:
                            r.text = r.text.replace(
                                'form.股东出资',
                                'company.股东出资'
                            )
                        if 'form.股东出资_历史' in r.text:
                            r.text = r.text.replace(
                                'form.股东出资_历史',
                                'company.股东出资_历史'
                            )

    doc.save('templates/FDD模板_循环版.docx')
    print(f'Loop added: start after P[{start_idx}], end before P[{end_idx}]')
    print(f'Saved: templates/FDD模板_循环版.docx')
else:
    print(f'Could not find boundaries. Headings: {heading_indices}, End: {end_indices}')
