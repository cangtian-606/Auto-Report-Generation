from docx import Document
doc = Document('templates/FDD初稿.docx')

# Show from 投资项目概况 to 项目的基本情况
show = False
stop = False
for i, p in enumerate(doc.paragraphs):
    style = p.style.name or ''
    text = p.text.strip()
    
    if '投资项目概况' in text and 'Heading 1' in style:
        show = True
    if '项目建设与运营' in text and 'Heading 1' in style:
        stop = True
    
    if not show or stop:
        continue
    if not text:
        # Show blank lines as separators
        continue
    
    if 'Heading' in style:
        level = style.replace('Heading ', '')
        prefix = '#' * int(level) if level.isdigit() else '##'
        print(f'P[{i}] {prefix} {text[:100]}')
    else:
        # Shorten very long texts
        display = text[:180] + ('...' if len(text) > 180 else '')
        # Mark 【】 and xxxx placeholders
        has_placeholder = '【' in display or 'xxxx' in display or '[]' in display
        mark = ' ⚠️' if has_placeholder else ''
        print(f'P[{i}] [{style}]{mark} {display}')

print()
print('=== TABLES in this range ===')
# Show tables and their relation to paragraphs
body = doc.element.body
para_elements = [p._p for p in doc.paragraphs]
table_elements = [t._tbl for t in doc.tables]

# Find table positions relative to paragraphs in the 投资项目概况-项目建设与运营 range
start_p_elem = None
end_p_elem = None
for i, p in enumerate(doc.paragraphs):
    if '投资项目概况' in p.text and 'Heading 1' in (p.style.name or ''):
        start_p_elem = p._p
    if '项目建设与运营' in p.text and 'Heading 1' in (p.style.name or ''):
        end_p_elem = p._p
        break

all_elems = list(body)
start_pos = all_elems.index(start_p_elem) if start_p_elem in all_elems else -1
end_pos = all_elems.index(end_p_elem) if end_p_elem in all_elems else len(all_elems)

for ti, t in enumerate(doc.tables):
    tbl_pos = all_elems.index(t._tbl) if t._tbl in all_elems else -1
    if start_pos <= tbl_pos <= end_pos:
        print(f'Table[{ti}] ({len(t.rows)}r x {len(t.columns)}c)')
        header = [str(c.text).replace('\n',' ')[:50] for c in t.rows[0].cells]
        print(f'  header: {header}')
        if len(t.rows) > 1:
            row1 = [str(c.text).replace('\n',' ')[:40] for c in t.rows[1].cells]
            print(f'  row1:   {row1}')
        if len(t.rows) > 2:
            rown = [str(c.text).replace('\n',' ')[:40] for c in t.rows[-1].cells]
            print(f'  last:   {rown}')
        print()
