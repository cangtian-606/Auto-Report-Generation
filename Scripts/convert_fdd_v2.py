"""将 FDD初稿.docx 转换为 docxtpl Jinja2 模板 (v2: 单次完成)"""

from pathlib import Path
from docx import Document

PROJECT_DIR = Path(__file__).resolve().parent.parent
SRC = PROJECT_DIR / "templates" / "FDD初稿.docx"
OUT = PROJECT_DIR / "templates" / "FDD模板.docx"


def set_cell_text(table, row, col, text):
    """Set cell text, creating runs if needed"""
    cell = table.cell(row, col)
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    p.clear()
    p.add_run(text)


def merge_cell_text(cell):
    """Get full text of a cell merged from all runs"""
    return ''.join(r.text for r in cell.paragraphs[0].runs) if cell.paragraphs[0].runs else ''


def replace_merge_cell(cell, old, new):
    """Merge runs, replace old with new, write back to first run"""
    full = ''
    for p in cell.paragraphs:
        for r in p.runs:
            full += r.text
    if old not in full:
        return False
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = full.replace(old, new)
    else:
        cell.paragraphs[0].add_run(full.replace(old, new))
    return True


def replace_para_text(para, old, new):
    """Replace text in all runs of a paragraph"""
    full = ''.join(r.text for r in para.runs)
    if old not in full:
        return False
    for r in para.runs:
        r.text = ''
    if para.runs:
        para.runs[0].text = full.replace(old, new)
    else:
        para.add_run(full.replace(old, new))
    return True


def convert():
    doc = Document(str(SRC))

    # ============ PARAGRAPHS ============
    for p in doc.paragraphs:
        text = ''.join(r.text for r in p.runs)
        if not text.strip():
            continue

        replace_para_text(p, '大华咨字[2026]XXXX号', '{{ date.全局.报告文号 }}')
        replace_para_text(p, '二〇二六年五月【】日', '二〇二六年五月{{ date.全局.报告日期_日 }}')

    # ============ TABLE[0]: Cover ============
    t0 = doc.tables[0]
    replace_merge_cell(t0.cell(0, 0), '【】', '{{ date.全局.项目公司简称 }}')
    replace_merge_cell(t0.cell(2, 0), '大华咨字[2026]XXXX号', '{{ date.全局.报告文号 }}')

    # ============ TABLE[1]: Company name mapping ============
    t1 = doc.tables[1]
    for ri in range(1, len(t1.rows)):
        replace_merge_cell(t1.cell(ri, 0), 'xxxx', '{{ date.全局.项目公司全称 }}')
        replace_merge_cell(t1.cell(ri, 1), 'xxxx', '{{ date.全局.项目公司简称 }}')

    # ============ TABLE[2]: Basic company info ============
    t2 = doc.tables[2]
    field_t2 = {
        '公司名称': '{{ date.基本情况.公司名称 }}',
        '统一社会信用代码': '{{ date.基本情况.信用代码 }}',
        '注册地址': '{{ date.基本情况.注册地址 }}',
        '法定代表人': '{{ date.基本情况.法定代表人 }}',
        '公司类型': '{{ date.基本情况.公司类型 }}',
        '注册资本': '{{ date.基本情况.注册资本 | num }}',
        '经营范围': '{{ date.基本情况.经营范围 }}',
        '成立日期': '{{ date.基本情况.成立日期 }}',
        '营业期限': '{{ date.基本情况.营业期限 }}',
    }
    for ri in range(len(t2.rows)):
        key = t2.cell(ri, 0).text.strip()
        if key in field_t2:
            replace_merge_cell(t2.cell(ri, 1), 'xxxx', field_t2[key])

    # ============ TABLE[3]: Shareholder → loop ============
    t3 = doc.tables[3]
    if len(t3.rows) >= 2:
        set_cell_text(t3, 1, 0, '{%tr for i in form.股东出资 %}')
        for ci in range(1, len(t3.columns)):
            set_cell_text(t3, 1, ci, '')
        if len(t3.rows) >= 3:
            set_cell_text(t3, 2, 0, '{{ i.股东名称 }}')
            set_cell_text(t3, 2, 1, '{{ i.认缴出资额 | money }}')
            set_cell_text(t3, 2, 2, '{{ i.认缴比例 | percent }}')
            set_cell_text(t3, 2, 3, '{{ i.实缴出资额 | money }}')
            if len(t3.columns) > 4:
                set_cell_text(t3, 2, 4, '{{ i.实缴比例 | percent }}')
        # Add endfor row if there's a 4th row, or replace last
        end_ri = 3 if len(t3.rows) >= 4 else len(t3.rows) - 1
        set_cell_text(t3, end_ri, 0, '{%tr endfor %}')
        for ci in range(1, len(t3.columns)):
            set_cell_text(t3, end_ri, ci, '')

    # ============ TABLE[4]: Blank basic info (second entity) ============
    t4 = doc.tables[4]
    for ri in range(len(t4.rows)):
        key = t4.cell(ri, 0).text.strip()
        if key in field_t2:
            replace_merge_cell(t4.cell(ri, 1), 'xxxx', field_t2[key])

    # ============ TABLE[5]: Historical shareholder → loop ============
    t5 = doc.tables[5]
    if len(t5.rows) >= 2:
        set_cell_text(t5, 1, 0, '{%tr for i in form.股东出资_历史 %}')
        set_cell_text(t5, 1, 1, '{{ i.股东名称 }}')
        set_cell_text(t5, 1, 2, '{{ i.认缴出资额 | money }}')
        set_cell_text(t5, 1, 3, '{{ i.认缴比例 | percent }}')
        if len(t5.columns) > 4:
            set_cell_text(t5, 1, 4, '{{ i.实缴出资额 | money }}')

    # ============ TABLE[6]: Project approvals ============
    t6 = doc.tables[6]
    replace_merge_cell(t6.cell(0, 2), '【】', '{{ date.全局.项目公司简称 }}')

    # ============ TABLE[7]: PC Contract cost ============
    t7 = doc.tables[7]
    set_cell_text(t7, 1, 0, '1')
    set_cell_text(t7, 2, 0, '2')
    set_cell_text(t7, 3, 0, '3')
    set_cell_text(t7, 4, 0, '合计')
    set_cell_text(t7, 4, 1, '')
    replace_merge_cell(t7.cell(1, 3), 'xxxx', '{{ date.PC合同.设备费 | money }}')
    replace_merge_cell(t7.cell(2, 3), 'xxxx', '{{ date.PC合同.建安费 | money }}')
    replace_merge_cell(t7.cell(3, 3), 'xxxx', '{{ date.PC合同.其他费 | money }}')
    replace_merge_cell(t7.cell(4, 3), 'xxxx', '{{ date.PC合同.合同总价 | money }}')
    replace_merge_cell(t7.cell(4, 1), 'xxxx', '')

    # ============ TABLE[8]: Tax rates → KEEP AS-IS (fixed) ============

    # ============ TABLE[9]: Fixed assets params ============
    t9 = doc.tables[9]
    replace_merge_cell(t9.cell(0, 2), '【】项目', '{{ date.全局.项目公司简称 }}项目')
    fa_fields = {
        '固定资产原值': '{{ date.固定资产.原值 | money }}',
        '残值率': '{{ date.固定资产.残值率 | percent }}',
        '折旧年限': '{{ date.固定资产.折旧年限 }}',
        '转固时点': '{{ date.固定资产.转固时点 }}',
        '年折旧额': '{{ date.固定资产.年折旧额 | money }}',
    }
    for ri in range(1, len(t9.rows)):
        key = t9.cell(ri, 0).text.strip()
        for k, v in fa_fields.items():
            if k in key:
                set_cell_text(t9, ri, 2, v)
                break

    # ============ TABLE[10]: Power generation ============
    t10 = doc.tables[10]
    period_names = ['2024年度', '2025年度', '2026年1至4月']
    for ri in range(1, min(len(t10.rows), len(period_names) + 1)):
        pn = period_names[ri - 1]
        set_cell_text(t10, ri, 0, pn)
        set_cell_text(t10, ri, 1, '{{ date.发电量.' + pn + '.自发自用 | num }}')
        set_cell_text(t10, ri, 2, '{{ date.发电量.' + pn + '.余电上网 | num }}')
        set_cell_text(t10, ri, 3, '{{ date.发电量.' + pn + '.发电量小计 | num }}')
        set_cell_text(t10, ri, 4, '{{ date.发电量.' + pn + '.自发自用占比 | percent }}')

    # ============ TABLE[11]: O&M contract ============
    t11 = doc.tables[11]
    om_fields = {
        '服务期限': '{{ date.运维合同.服务期限 }}',
        '合同金额（含税）': '{{ date.运维合同.合同金额 | money }}',
        '合同单价（含税）': '{{ date.运维合同.合同单价 | money }}',
        '进项税率': '{{ date.运维合同.进项税率 }}',
    }
    for ri in range(len(t11.rows)):
        key = t11.cell(ri, 0).text.strip()
        if key in om_fields:
            set_cell_text(t11, ri, 1, om_fields[key])

    # ============ TABLE[12]: Balance Sheet ============
    t12 = doc.tables[12]
    period_keys = ['20260430', '20251231', '20241231']
    for ri in range(1, len(t12.rows)):
        field = t12.cell(ri, 0).text.strip()
        if field:
            safe = field.replace(' ', '_').replace('（', '(').replace('）', ')').replace('/', '_')
            for ci in range(1, min(len(t12.columns), len(period_keys) + 1)):
                set_cell_text(t12, ri, ci, '{{ date.资产负债表.' + safe + '_' + period_keys[ci - 1] + ' | money }}')

    # ============ TABLE[13]: Income Statement ============
    t13 = doc.tables[13]
    pl_periods = ['2026年1至4月', '2025年度', '2024年度']
    for ri in range(1, len(t13.rows)):
        field = t13.cell(ri, 0).text.strip()
        if field:
            safe = field.replace(' ', '_').replace('（', '(').replace('）', ')')
            for ci in range(1, min(len(t13.columns), len(pl_periods) + 1)):
                set_cell_text(t13, ri, ci, '{{ date.利润表.' + safe + '_' + pl_periods[ci - 1].replace('-', '') + ' | money }}')

    # ============ TABLE[14]: Related parties → loop ============
    t14 = doc.tables[14]
    if len(t14.rows) > 1:
        # Remove all rows except header
        for ri in range(len(t14.rows) - 1, 0, -1):
            t14.rows[ri]._tr.getparent().remove(t14.rows[ri]._tr)
        r1 = t14.add_row()
        set_cell_text(t14, 1, 0, '{%tr for i in form.关联方 %}')
        set_cell_text(t14, 1, 1, '{{ i.关联方名称 }}')
        set_cell_text(t14, 1, 2, '{{ i.关联关系 }}')
        r2 = t14.add_row()
        set_cell_text(t14, 2, 0, '{%tr endfor %}')
        set_cell_text(t14, 2, 1, '')
        set_cell_text(t14, 2, 2, '')

    # ============ TABLE[15]: Related party balances → loop ============
    t15 = doc.tables[15]
    if len(t15.rows) > 1:
        for ri in range(len(t15.rows) - 1, 0, -1):
            t15.rows[ri]._tr.getparent().remove(t15.rows[ri]._tr)
        r1 = t15.add_row()
        set_cell_text(t15, 1, 0, '{%tr for i in form.关联方余额 %}')
        set_cell_text(t15, 1, 1, '{{ i.关联方 }}')
        set_cell_text(t15, 1, 2, '{{ i.余额_20260430 | money }}')
        set_cell_text(t15, 1, 3, '{{ i.余额_20251231 | money }}')
        set_cell_text(t15, 1, 4, '{{ i.余额_20241231 | money }}')
        r2 = t15.add_row()
        set_cell_text(t15, 2, 0, '{%tr endfor %}')
        for ci in range(1, 5):
            set_cell_text(t15, 2, ci, '')

    # ============ TABLE[16]: Related party transactions → loop ============
    t16 = doc.tables[16]
    if len(t16.rows) > 1:
        for ri in range(len(t16.rows) - 1, 0, -1):
            t16.rows[ri]._tr.getparent().remove(t16.rows[ri]._tr)
        r1 = t16.add_row()
        set_cell_text(t16, 1, 0, '{%tr for i in form.关联方交易 %}')
        set_cell_text(t16, 1, 1, '{{ i.关联方 }}')
        set_cell_text(t16, 1, 2, '{{ i.金额_2026年1至4月 | money }}')
        set_cell_text(t16, 1, 3, '{{ i.金额_2025年度 | money }}')
        set_cell_text(t16, 1, 4, '{{ i.金额_2024年度 | money }}')
        r2 = t16.add_row()
        set_cell_text(t16, 2, 0, '{%tr endfor %}')
        for ci in range(1, 5):
            set_cell_text(t16, 2, ci, '')

    doc.save(str(OUT))
    print(f"Template saved: {OUT}")


if __name__ == '__main__':
    convert()
