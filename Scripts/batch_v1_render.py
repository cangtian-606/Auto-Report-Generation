"""Generate stamped company A & B xlsx files from yaml, then render both"""

import sys; sys.path.insert(0, '.')
from pathlib import Path
from src.yaml_reader import YamlDataReader
from src.generator import DocumentGenerator
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

PROJECT = Path(__file__).resolve().parent.parent

for company in ["A公司", "B公司"]:
    print("=" * 60)
    print("  %s" % company)
    print("=" * 60)

    # 1. Read yaml
    reader = YamlDataReader(str(PROJECT / "data" / ("FDDv1_%s.yaml" % company)))
    ctx = reader.read_context()

    # 2. Write xlsx
    wb = openpyxl.Workbook()
    hf = Font(bold=True, size=11)
    hfill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")

    def write_kv(ws, data, sheet_name):
        ws.title = sheet_name
        ws.column_dimensions['A'].width = 28
        ws.column_dimensions['B'].width = 55
        ws.append(["字段编码", "值"])
        ws['A1'].font = hf; ws['B1'].font = hf
        ws['A1'].fill = hfill; ws['B1'].fill = hfill
        for k, v in data.items():
            ws.append([k, v])

    def write_table(ws, rows, columns, sheet_name):
        ws.title = sheet_name
        for ci, col in enumerate(columns, 1):
            c = ws.cell(row=1, column=ci, value=col)
            c.font = hf; c.fill = hfill
        ws.column_dimensions['A'].width = 22
        ws.column_dimensions['B'].width = 22
        for ri, row in enumerate(rows, 2):
            for ci, col in enumerate(columns, 1):
                ws.cell(row=ri, column=ci, value=row.get(col, ""))

    # date sheets
    first = True
    for key, data in ctx['date'].items():
        if first:
            write_kv(wb.active, data, "date.%s" % key)
            first = False
        else:
            ws = wb.create_sheet()
            write_kv(ws, data, "date.%s" % key)

    # form sheets -also write empty tables with header columns
    form_columns = {
        "释义": ["全称", "简称"],
        "股权结构": ["股东名称", "认缴出资额", "认缴比例", "实缴出资额", "实缴比例"],
        "股东信息": ["股东名称", "股东介绍"],
        "公司设立": ["股东名称", "认缴出资额", "认缴比例", "实缴出资额", "实缴比例"],
    }
    for key in form_columns:
        rows = ctx['form'].get(key, [])
        ws = wb.create_sheet("form.%s" % key)
        columns = form_columns[key]
        write_table(ws, rows, columns, "form.%s" % key)

    xlsx_path = PROJECT / "data" / ("%s.xlsx" % company)
    wb.save(str(xlsx_path))
    print("  XLSX: %s" % xlsx_path)

    # 3. Render
    gen = DocumentGenerator(str(PROJECT / "templates" / "FDD初稿 v1.docx"))
    out = PROJECT / "output" / ("%s_output.docx" % company)
    ok = gen.render(ctx, str(out))
    print("  Render: %s" % ("OK" if ok else "FAIL"))

    # Verify
    if ok:
        from docx import Document
        doc = Document(str(out))
        all_text = ''
        for p in doc.paragraphs: all_text += p.text + '\n'
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells: all_text += c.text + '\n'

        has_tags = '{{' in all_text
        has_company_name = ctx['date']['全局']['公司简称'] in all_text
        print("  Unrendered tags: %s" % has_tags)
        print("  Company name in output: %s" % has_company_name)

        # Check a few key replacements
        checks = [
            ctx['date']['全局']['项目名称'],
            ctx['date']['历史沿革']['统一社会信用代码'],
        ]
        if ctx['form']['股权结构']:
            checks.append(ctx['form']['股权结构'][0]['股东名称'])
        if ctx['form']['公司设立']:
            checks.append(ctx['form']['公司设立'][0]['股东名称'])
        for c in checks:
            ok_flag = c in all_text
            print("    [%s]: %s" % (c[:30], "OK" if ok_flag else "MISSING"))

print("\nDone. Both companies rendered.")
