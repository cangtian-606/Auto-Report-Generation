"""Verify nested data end-to-end with a simple template"""
from docx import Document

doc = Document()

doc.add_paragraph('{% for company in form.项目公司 %}')
doc.add_paragraph()

doc.add_paragraph('================================================================')
doc.add_paragraph('  {{ company.公司简称 }} 尽职调查报告')
doc.add_paragraph('================================================================')
doc.add_paragraph()
doc.add_paragraph('公司名称：{{ company.公司名称 }}')
doc.add_paragraph('信用代码：{{ company.信用代码 }}')
doc.add_paragraph('注册资本：{{ company.注册资本 | num }}万元')
doc.add_paragraph('成立日期：{{ company.成立日期 }}')
doc.add_paragraph()
doc.add_paragraph('股东出资情况：')

table = doc.add_table(rows=4, cols=5)
table.style = 'Table Grid'
headers = ['股东', '认缴出资额', '认缴比例', '实缴出资额', '实缴比例']
for ci, h in enumerate(headers):
    table.cell(0, ci).text = h

table.cell(1, 0).paragraphs[0].add_run('{%tr for i in company.股东出资 %}')
table.cell(2, 0).paragraphs[0].add_run('{{ i.股东 }}')
table.cell(2, 1).paragraphs[0].add_run('{{ i.认缴出资额 | money }}')
table.cell(2, 2).paragraphs[0].add_run('{{ i.认缴比例 | percent }}')
table.cell(2, 3).paragraphs[0].add_run('{{ i.实缴出资额 | money }}')
table.cell(2, 4).paragraphs[0].add_run('{{ i.实缴比例 | percent }}')
table.cell(3, 0).paragraphs[0].add_run('{%tr endfor %}')

doc.add_paragraph()
doc.add_paragraph('{% endfor %}')
doc.add_paragraph()
doc.add_paragraph('—— 报告结束 ——')

doc.save('templates/嵌套测试模板.docx')
print('Saved: templates/嵌套测试模板.docx')
