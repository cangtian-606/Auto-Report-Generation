"""修复尽调报告测试数据并验证渲染
根因：项目基本情况 Sheet 被读取为 table (list)，
但模板用 {% for i in 项目基本情况.项目公司 %} 期望 dict.项目公司 (list)。

修复：将项目基本情况拆分为 KV Sheet + 子表 Sheet。
"""
import sys, os
from pathlib import Path
project_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(project_root))

from openpyxl import Workbook
from src.orchestrator import generate

data_path = project_root / 'data' / '尽调报告_测试数据_v2.xlsx'
template_path = project_root / 'templates' / '尽调报告模板v1.docx'
output_path = project_root / 'output' / '尽调报告_测试输出.docx'

# ===== 创建正确结构的 xlsx =====
wb = Workbook()
wb.remove(wb.active)

# Sheet 1: 全局 (KV)
ws1 = wb.create_sheet('全局')
ws1.append(['字段编码', '值'])
ws1.append(['公司名称', '重庆晟和泰新能源科技有限公司'])
ws1.append(['公司简称', '重庆晟和泰'])
ws1.append(['报告文号', '2025-001'])

# Sheet 2: 历史沿革 (KV)
ws2 = wb.create_sheet('历史沿革')
ws2.append(['字段编码', '值'])
ws2.append(['公司名称', '重庆晟和泰新能源科技有限公司'])
ws2.append(['统一社会信用代码', '91500227MA5U7XXXXX'])
ws2.append(['注册地址', '重庆市巫山县工业园区'])
ws2.append(['法定代表人', '张三'])
ws2.append(['公司类型', '有限责任公司(非自然人投资或控股的法人独资)'])
ws2.append(['注册资本', 10000000])
ws2.append(['经营范围', '太阳能发电、光伏发电项目的开发、建设、运营及管理'])
ws2.append(['成立日期', '2018-06-15'])
ws2.append(['营业期限', '2018-06-15至无固定期限'])
ws2.append(['设立公司', '成都交投新能源集团有限公司'])
ws2.append(['设立日期', '2018-06-15'])
ws2.append(['设立工商局', '重庆市巫山县市场监督管理局'])

# Sheet 3: 历史沿革.股权结构 (表格 → 键值对父表的子表)
ws3 = wb.create_sheet('历史沿革.股权结构')
ws3.append(['_parent_公司名称', '股东名称', '认缴出资额', '认缴比例', '实缴出资额', '实缴比例'])
ws3.append(['重庆晟和泰新能源科技有限公司', '成都交投新能源集团有限公司', 10000000, 1.0, 10000000, 1.0])

# Sheet 4: 历史沿革.公司设立 (表格 → 键值对父表的子表)
ws4 = wb.create_sheet('历史沿革.公司设立')
ws4.append(['_parent_公司名称', '股东名称', '认缴出资额', '认缴比例', '实缴出资额', '实缴比例'])
ws4.append(['重庆晟和泰新能源科技有限公司', '成都交投新能源集团有限公司', 10000000, 1.0, 10000000, 1.0])

# Sheet 5: 释义 (表格)
ws5 = wb.create_sheet('释义')
ws5.append(['全称', '简称'])
ws5.append(['成都交投新能源集团有限公司', '成都交投'])
ws5.append(['重庆晟和泰新能源科技有限公司', '重庆晟和泰'])
ws5.append(['财务尽职调查', '尽调'])

# ===== 关键修复 =====
# Sheet 6: 项目基本情况 → 改为 KV (dict)，不再是表格
ws6 = wb.create_sheet('项目基本情况')
ws6.append(['字段编码', '值'])
# 空 KV Sheet 即可——它只需要占位，使得 context["项目基本情况"] 是一个 dict

# Sheet 7: 项目基本情况.项目公司 → 表格 (list)，挂载到父 dict 的 "项目公司" key
ws7 = wb.create_sheet('项目基本情况.项目公司')
ws7.append(['项目名称', '项目地址', '项目装机量', '项目模式', '含税总价', '含税单价'])
ws7.append(['重庆晟和泰分布式光伏一期', '重庆市巫山县工业园区A区', '5MW', '全额上网', 20000000, 4.00])
ws7.append(['重庆晟和泰分布式光伏二期', '重庆市巫山县工业园区B区', '3MW', '自发自用', 12600000, 4.20])

wb.save(str(data_path))
print(f'Data: {data_path}')
print(f'Sheets: {wb.sheetnames}')

# ===== 渲染 =====
print()
success = generate(
    str(data_path),
    str(template_path),
    str(output_path),
)
print(f'Render: {success}')

if success:
    from docx import Document
    doc = Document(str(output_path))
    # 检查关键段落
    print('\n=== 项目概况段落验证 ===')
    found_1 = found_2 = found_3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if '重庆晟和泰分布式光伏' in t:
            print(f'  [OK] {t[:120]}')
            found_1 = True
        if '20,000,000.00' in t:
            print(f'  [OK] {t[:120]}')
            found_2 = True
        if '4.00' in t and 'Wp' in t:
            print(f'  [OK] {t[:120]}')
            found_3 = True

    print(f'\n项目概况: {"PASS" if found_1 else "FAIL"}')
    print(f'项目总投资: {"PASS" if found_2 else "FAIL"}')
    print(f'项目单价: {"PASS" if found_3 else "FAIL"}')

    # Check tc loop in table 5
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            cells = [c.text.strip()[:40] for c in row.cells]
            combined = ' | '.join(cells)
            if '分布式光伏' in combined or '5MW' in combined:
                print(f'\nTable{ti} R{ri}: {combined[:150]}')
