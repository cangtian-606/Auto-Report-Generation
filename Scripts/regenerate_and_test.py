"""
v3: 重建手续表数据结构 — 手续字段直接挂到项目 dict 上
模板用 {{ p.手续_R01 }} 替代 {{ 手续.R01[p.项目名称] }}
"""
import sys
from pathlib import Path
project = Path(r'D:\ProgramWorkingSpace\Auto-Report-Generation')
sys.path.insert(0, str(project))

from openpyxl import Workbook
from docx import Document
from src.orchestrator import generate

# ============================================================
# 1. 修复模板 — TC 循环变量改用 p.手续_RXX
# ============================================================
tpl_path = project / 'templates' / '尽调报告模板v1.docx'
doc = Document(str(tpl_path))
tbl = doc.tables[5]

for ri in range(1, len(tbl.rows)):
    row = tbl.rows[ri]
    # cell5 清空
    if len(row.cells) > 5:
        row.cells[5].text = ''
    # cell3 改写: {{ p.手续_RXX }}
    row.cells[3].text = ''
    row.cells[3].paragraphs[0].add_run(
        '{{ p.手续_R' + str(ri).zfill(2) + ' }}'
    )

doc.save(str(tpl_path))
print("Template tc vars fixed: {{ p.手续_RXX }}")

# ============================================================
# 2. 生成 xlsx 测试数据
# ============================================================
data_path = project / 'data' / '尽调报告_测试数据_v3.xlsx'
wb = Workbook()
wb.remove(wb.active)

def add_kv_sheet(wb, name, pairs):
    ws = wb.create_sheet(name)
    ws.append(['字段编码', '值'])
    for k, v in pairs:
        ws.append([k, v])

def add_table_sheet(wb, name, headers, rows):
    ws = wb.create_sheet(name)
    ws.append(headers)
    for row in rows:
        ws.append(row)

add_kv_sheet(wb, '全局', [
    ('公司名称', '重庆晟和泰新能源科技有限公司'),
    ('公司简称', '重庆晟和泰'),
    ('报告文号', '2025-001'),
])

add_kv_sheet(wb, '历史沿革', [
    ('公司名称', '重庆晟和泰新能源科技有限公司'),
    ('统一社会信用代码', '91500227MA5U7XXXXX'),
    ('注册地址', '重庆市巫山县工业园区'),
    ('法定代表人', '张三'),
    ('公司类型', '有限责任公司(非自然人投资或控股的法人独资)'),
    ('注册资本', 10000000),
    ('经营范围', '太阳能发电、光伏发电项目的开发、建设、运营及管理'),
    ('成立日期', '2018-06-15'),
    ('营业期限', '2018-06-15至无固定期限'),
    ('设立公司', '成都交投新能源集团有限公司'),
    ('设立日期', '2018-06-15'),
    ('设立工商局', '重庆市巫山县市场监督管理局'),
])

add_table_sheet(wb, '历史沿革.股权结构',
    ['_parent_公司名称', '股东名称', '认缴出资额', '认缴比例', '实缴出资额', '实缴比例'],
    [['重庆晟和泰新能源科技有限公司', '成都交投新能源集团有限公司', 10000000, 1.0, 10000000, 1.0],
     ])

add_table_sheet(wb, '历史沿革.公司设立',
    ['_parent_公司名称', '股东名称', '认缴出资额', '认缴比例', '实缴出资额', '实缴比例'],
    [['重庆晟和泰新能源科技有限公司', '成都交投新能源集团有限公司', 10000000, 1.0, 10000000, 1.0],
     ])

add_table_sheet(wb, '释义',
    ['全称', '简称'],
    [['成都交投新能源集团有限公司', '成都交投'],
     ['重庆晟和泰新能源科技有限公司', '重庆晟和泰'],
     ['财务尽职调查', '尽调'],
     ])

# 项目基本情况: KV 占位 → context["项目基本情况"] = {}
add_kv_sheet(wb, '项目基本情况', [])

# 项目公司: 完整列表
add_table_sheet(wb, '项目基本情况.项目公司',
    ['项目名称', '项目地址', '项目装机量', '项目模式', '含税总价', '含税单价'],
    [['重庆晟和泰分布式光伏一期', '重庆市巫山县工业园区A区', '5MW', '全额上网', 20000000, 4.00],
     ['重庆晟和泰分布式光伏二期', '重庆市巫山县工业园区B区', '3MW', '自发自用', 12600000, 4.20],
     ])

# 项目公司_手续: 预过滤 + 手续字段直接挂在 dict 上
add_table_sheet(wb, '项目基本情况.项目公司_手续',
    ['项目名称', '项目地址', '项目装机量', '项目模式', '含税总价', '含税单价',
     '手续_R01', '手续_R02', '手续_R03', '手续_R04',
     '手续_R05', '手续_R06', '手续_R07', '手续_R08',
     '手续_R09', '手续_R10', '手续_R11', '手续_R12'],
    [['重庆晟和泰分布式光伏一期', '重庆市巫山县工业园区A区', '5MW', '全额上网', 20000000, 4.00,
      '是', '是', '是', '是', '', '是', '是', '是', '是', '', '是', '是'],
     ])

wb.save(str(data_path))
print("Data saved: " + str(data_path))
print("Sheets: " + str(wb.sheetnames))

# ============================================================
# 3. 渲染 + 验证
# ============================================================
print()
out_path = project / 'output' / '尽调报告_测试输出.docx'
success = generate(str(data_path), str(tpl_path), str(out_path))
print("Render: " + ("PASS" if success else "FAIL"))

if success:
    out_doc = Document(str(out_path))
    print("\n===== 输出验证 =====")

    # 1. 查 for 块无空行
    print("--- 项目概况 (for 块不应有空行) ---")
    count = 0
    for p in out_doc.paragraphs:
        t = p.text.strip()
        if '分布式光伏' in t and ('位于' in t or '装机容量' in t):
            count += 1
            print("  [" + str(count) + "] " + t[:130])

    # 2. 查项目总投资
    print("\n--- 项目总投资 ---")
    for p in out_doc.paragraphs:
        t = p.text.strip()
        if '含税工程总价' in t or '含税建设单价' in t:
            print("  " + t[:130])

    # 3. 查手续表
    print("\n--- 手续表 R00-R02 ---")
    tbl5 = out_doc.tables[5]
    for ri in range(min(3, len(tbl5.rows))):
        cells = [c.text.strip()[:35] for c in tbl5.rows[ri].cells]
        print("  R{:02d}: {}".format(ri, " | ".join(cells)))

    # 4. 统计是否有未渲染变量
    bad = False
    for p in out_doc.paragraphs:
        if '{{' in p.text and '}}' in p.text:
            print("  UNRENDERED: " + p.text[:100])
            bad = True
    for tbl in out_doc.tables:
        for r in tbl.rows:
            for c in r.cells:
                if '{{' in c.text and '}}' in c.text:
                    print("  UNRENDERED cell: " + c.text[:80])
                    bad = True
    if not bad:
        print("  All variables rendered!")

    print("\nSize: {:,} bytes".format(out_path.stat().st_size))
