"""Fix remaining placeholders in FDD template"""
from docx import Document

doc = Document('templates/FDD_项目概况模板.docx')

def set_cell_text(table, row, col, text):
    cell = table.cell(row, col)
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    cell.paragraphs[0].clear()
    cell.paragraphs[0].add_run(text)

# ============ PARAGRAPHS: remaining 【】 ============
for i, p in enumerate(doc.paragraphs):
    full = ''.join(r.text for r in p.runs)

    # P[65]: 投资【】 → 投资{{ 全局.公司简称 }}
    if '投资【】规划时参考' in full:
        for r in p.runs: r.text = ''
        new = full.replace('投资【】规划时参考', '投资{{ 全局.公司简称 }}规划时参考')
        if p.runs: p.runs[0].text = new

    # P[67]: 【】财务状况
    elif '截至2026年4月30日【】财务状况' in full:
        for r in p.runs: r.text = ''
        new = full.replace('【】财务状况', '{{ 全局.公司简称 }}财务状况')
        if p.runs: p.runs[0].text = new

    # P[69]: 就【】历史沿革
    elif '就【】历史沿革' in full:
        for r in p.runs: r.text = ''
        new = full.replace('就【】', '就{{ 全局.公司简称 }}')
        if p.runs: p.runs[0].text = new

    # P[73]: 【】的财务状况
    elif '截至2026年4月30日【】的财务状况' in full:
        for r in p.runs: r.text = ''
        new = full.replace('【】的财务状况', '{{ 全局.公司简称 }}的财务状况')
        if p.runs: p.runs[0].text = new

    # P[117]: xxxx → 工程承包范围 (before EMC合同 section)
    if i == 117 and full.strip() == 'xxxx':
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = '{{ PC合同.工程承包范围 }}'

    # P[124]: xxxx → 电费计算方式 (under EMC合同 / ④电费计算)
    if i == 124 and full.strip() == 'xxxx':
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = '{{ EMC合同.电费计算方式 }}'

    # P[126]: ⑤结算方式： [] → already handled but check
    elif '结算方式' in full and '[]' in full:
        for r in p.runs: r.text = ''
        new = full.replace('[]', '{{ EMC合同.结算方式 }}')
        if p.runs: p.runs[0].text = new

    # P[142]: 【】项目 → {{ 全局.公司简称 }}项目
    elif full.startswith('【】项目'):
        for r in p.runs: r.text = ''
        new = full.replace('【】', '{{ 全局.公司简称 }}')
        if p.runs: p.runs[0].text = new

    # P[178]: 二〇二六年五月【】日
    elif '二〇二六年五月【】日' in full:
        for r in p.runs: r.text = ''
        new = full.replace('二〇二六年五月【】日', '二〇二六年五月{{ 全局.报告日期_日 }}日')
        if p.runs: p.runs[0].text = new

# ============ TABLE FIXES ============
# T[4]: 第二实体基本信息 - clean xxxx
t4 = doc.tables[4]
field_map = {
    '统一社会信用代码': '{{ 基本情况.信用代码 }}',
    '注册地址': '{{ 基本情况.注册地址 }}',
    '法定代表人': '{{ 基本情况.法定代表人 }}',
    '公司类型': '{{ 基本情况.公司类型 }}',
    '注册资本': '{{ 基本情况.注册资本 | num }}',
    '经营范围': '{{ 基本情况.经营范围 }}',
    '成立日期': '{{ 基本情况.成立日期 }}',
    '营业期限': '{{ 基本情况.营业期限 }}',
}
for ri in range(len(t4.rows)):
    key = t4.cell(ri, 0).text.strip()
    if key in field_map:
        set_cell_text(t4, ri, 1, field_map[key])
    elif not key:
        set_cell_text(t4, ri, 0, '')
        set_cell_text(t4, ri, 1, '')

# T[6] R[0] C[2]: 【】 → {{ 全局.公司简称 }}
t6 = doc.tables[6]
for p in t6.cell(0, 2).paragraphs:
    full = ''.join(r.text for r in p.runs)
    if '【】' in full:
        for r in p.runs: r.text = ''
        if p.runs:
            p.runs[0].text = full.replace('【】', '{{ 全局.公司简称 }}')

# T[7]: 进项税率列 (col 2) - fill with fixed tax rates
t7 = doc.tables[7]
tax_rates = ['13%', '9%', '6%', '']
for ri in range(1, len(t7.rows)):
    rate = tax_rates[ri-1] if ri-1 < len(tax_rates) else ''
    set_cell_text(t7, ri, 2, rate)

doc.save('templates/FDD_项目概况模板.docx')

# Verify
doc2 = Document('templates/FDD_项目概况模板.docx')
all_text = ''
for p in doc2.paragraphs:
    all_text += ''.join(r.text for r in p.runs) + '\n'
for t in doc2.tables:
    for r in t.rows:
        for c in r.cells:
            all_text += c.text + '\n'

remaining = [c for c in ['【】', 'xxxx'] if c in all_text]
# [] are used in jinja2 [] bracket syntax, ignore those
bracket_only = all_text.count('[]') - all_text.count("['")
print(f'Remaining 【】/xxxx: {remaining}')
print(f'Jinja2 tags: {all_text.count("{{")}')
print(f'Loops: {all_text.count("{%tr")}')
