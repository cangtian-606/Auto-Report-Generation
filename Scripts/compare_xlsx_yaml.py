"""Compare XLSX vs YAML render output for the same company"""
import sys; sys.path.insert(0, '.')
from src.reader import ExcelDataReader
from src.mapper import DataMapper
from src.yaml_reader import YamlDataReader
from src.generator import DocumentGenerator
from docx import Document

def doc_to_text(doc):
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append('[P] ' + t[:150])
    for ti, t in enumerate(doc.tables):
        for ri, r in enumerate(t.rows):
            cells = [c.text.strip()[:50] for c in r.cells]
            if any(c for c in cells):
                lines.append('[T%d R%d] %s' % (ti, ri, ' | '.join(cells)))
    return '\n'.join(lines)

def compare_for_company(label, yaml_file, xlsx_file):
    print("=" * 70)
    print("  %s" % label)
    print("=" * 70)

    # YAML path
    reader_yaml = YamlDataReader(yaml_file)
    ctx_yaml = reader_yaml.read_context()

    # XLSX path
    reader_xlsx = ExcelDataReader(xlsx_file)
    raw_xlsx = reader_xlsx.read_all()
    mapper = DataMapper(raw_xlsx)
    ctx_xlsx = mapper.build_context()

    print("\n--- context differences ---")

    # Compare date keys
    date_keys_y = set(ctx_yaml['date'].keys())
    date_keys_x = set(ctx_xlsx['date'].keys())
    only_yaml = date_keys_y - date_keys_x
    only_xlsx = date_keys_x - date_keys_y
    common = date_keys_y & date_keys_x
    if only_yaml: print("  date keys only in YAML: %s" % only_yaml)
    if only_xlsx: print("  date keys only in XLSX: %s" % only_xlsx)

    for key in sorted(common):
        yv = ctx_yaml['date'][key]
        xv = ctx_xlsx['date'][key]
        y_keys = set(yv.keys())
        x_keys = set(xv.keys())
        if y_keys != x_keys:
            print("  date.%s keys differ:" % key)
            print("    YAML extra: %s" % (y_keys - x_keys))
            print("    XLSX extra: %s" % (x_keys - y_keys))
        else:
            for fk in y_keys:
                yf = yv[fk]
                xf = xv[fk]
                if yf != xf:
                    print("  date.%s.%s: YAML=%r  XLSX=%r" % (key, fk, yf, xf))

    # Compare form keys
    form_keys_y = set(ctx_yaml['form'].keys())
    form_keys_x = set(ctx_xlsx['form'].keys())
    print("\n  form keys YAML: %s (%d)" % (form_keys_y, len(form_keys_y)))
    print("  form keys XLSX: %s (%d)" % (form_keys_x, len(form_keys_x)))

    for key in sorted(form_keys_y | form_keys_x):
        y_rows = ctx_yaml['form'].get(key, [])
        x_rows = ctx_xlsx['form'].get(key, [])
        if len(y_rows) != len(x_rows):
            print("  form.%s: rows YAML=%d XLSX=%d" % (key, len(y_rows), len(x_rows)))
        elif y_rows and x_rows:
            y_cols = set(y_rows[0].keys())
            x_cols = set(x_rows[0].keys())
            if y_cols != x_cols:
                print("  form.%s cols differ: YAML=%s XLSX=%s" % (key, y_cols, x_cols))
            else:
                for ri, (yr, xr) in enumerate(zip(y_rows, x_rows)):
                    for k in y_cols:
                        yv = yr[k]
                        xv = xr[k]
                        if yv != xv:
                            if k == '认缴出资额' or k == '实缴出资额' or k == '注册资本':
                                continue
                            print("  form.%s[%d].%s: YAML=%r  XLSX=%r" % (key, ri, k, yv, xv))

    # Render both
    print("\n--- rendering ---")
    gen_y = DocumentGenerator('templates/FDD初稿 v1.docx')
    gen_x = DocumentGenerator('templates/FDD初稿 v1.docx')
    
    gen_y.render(ctx_yaml, 'output/_cmp_yaml.docx')
    gen_x.render(ctx_xlsx, 'output/_cmp_xlsx.docx')

    doc_y = Document('output/_cmp_yaml.docx')
    doc_x = Document('output/_cmp_xlsx.docx')

    text_y = doc_to_text(doc_y)
    text_x = doc_to_text(doc_x)

    # Line by line diff
    lines_y = text_y.split('\n')
    lines_x = text_x.split('\n')
    
    diffs = 0
    max_lines = max(len(lines_y), len(lines_x))
    for i in range(max_lines):
        ly = lines_y[i] if i < len(lines_y) else '(MISSING)'
        lx = lines_x[i] if i < len(lines_x) else '(MISSING)'
        if ly != lx:
            diffs += 1
            if diffs <= 30:
                print("  L%d diff:" % i)
                print("    YAML: %s" % ly[:160])
                print("    XLSX: %s" % lx[:160])

    if diffs == 0:
        print("  IDENTICAL OUTPUT")
    else:
        print("  Total differences: %d lines" % diffs)


compare_for_company(
    "A公司 (安徽富军)",
    "data/FDDv1_A公司.yaml",
    "data/A公司.xlsx"
)

print()
print()
compare_for_company(
    "B公司 (重庆晟和泰)",
    "data/FDDv1_B公司.yaml",
    "data/B公司.xlsx"
)
