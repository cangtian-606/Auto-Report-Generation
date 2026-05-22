"""FDD初稿 → 模板转换脚本（封面~项目的基本情况）"""

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import re

PROJECT_DIR = Path(__file__).resolve().parent.parent
SRC = PROJECT_DIR / "templates" / "FDD初稿.docx"
OUT = PROJECT_DIR / "templates" / "FDD_项目概况模板.docx"


def set_cell_text(table, row, col, text):
    cell = table.cell(row, col)
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    cell.paragraphs[0].clear()
    cell.paragraphs[0].add_run(text)


def replace_para_text(para, old, new):
    full = ''.join(r.text for r in para.runs)
    if old not in full:
        return False
    for r in para.runs:
        r.text = ''
    if para.runs:
        para.runs[0].text = full.replace(old, new)
    return True


def replace_header_text(header, old, new):
    """Replace text in header paragraphs"""
    for p in header.paragraphs:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)


def convert():
    doc = Document(str(SRC))
    print(f"Loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {len(doc.sections)} sections")

    # ============ SECTION 0: COVER ============
    t0 = doc.tables[0]
    # Row 0 col 0: 【】 → 公司简称
    for p in t0.cell(0, 0).paragraphs:
        full = ''.join(r.text for r in p.runs)
        if '【】' in full:
            for r in p.runs: r.text = ''
            if p.runs: p.runs[0].text = full.replace('【】', '{{ 全局.公司简称 }}')
    # Row 2 col 0: 文号
    for p in t0.cell(2, 0).paragraphs:
        full = ''.join(r.text for r in p.runs)
        for r in p.runs: r.text = ''
        if p.runs:
            p.runs[0].text = '大华咨字[{{ 全局.报告文号 }}]号'
    print("[OK] Cover page")

    # ============ SECTION 1: 释义 (Table[1]) - form.用语释义 ============
    t1 = doc.tables[1]
    # row0 = header (全称 | 简称), rows 1-4 are data
    for ri in range(1, len(t1.rows)):
        replace_para_text(t1.cell(ri, 0).paragraphs[0], 'xxxx', '{{ i.全称 }}')
        replace_para_text(t1.cell(ri, 1).paragraphs[0], 'xxxx', '{{ i.简称 }}')
    print("[OK] Table[1] 用语释义")

    # ============ PARAGRAPHS: global replacements ============
    for p in doc.paragraphs:
        # Replace 2026年4月30日 with base date variable
        for r in p.runs:
            r.text = r.text.replace('2026年4月30日', '{{ 全局.尽调基准日 }}')

    print("[OK] Date replacements in paragraphs")

    # ============ SECTION 2: HEADERS ============
    for si, section in enumerate(doc.sections):
        header = section.header
        if header:
            for p in header.paragraphs:
                for r in p.runs:
                    if '大华咨字[2026]XXXX号' in r.text:
                        r.text = r.text.replace(
                            '大华咨字[2026]XXXX号',
                            '{{ 全局.报告文号 }}'
                        )
                        print(f"[OK] Header section[{si}]")

    # ============ TABLE[2]: 基本信息 (9 rows key-value) ============
    t2 = doc.tables[2]
    # row[0]: 公司名称 | xxxx  → 公司名称 | {{ 基本情况.公司名称 }}
    # row[1]: 统一社会信用代码 | xxxx → 统一社会信用代码 | {{ 基本情况.信用代码 }}
    # row[2]: 注册地址 | xxxx → ...
    # row[8]: 营业期限 | xxxx → ...
    field_map_t2 = {
        '公司名称': '{{ 基本情况.公司名称 }}',
        '统一社会信用代码': '{{ 基本情况.信用代码 }}',
        '注册地址': '{{ 基本情况.注册地址 }}',
        '法定代表人': '{{ 基本情况.法定代表人 }}',
        '公司类型': '{{ 基本情况.公司类型 }}',
        '注册资本': '{{ 基本情况.注册资本 | num }}',
        '经营范围': '{{ 基本情况.经营范围 }}',
        '成立日期': '{{ 基本情况.成立日期 }}',
        '营业期限': '{{ 基本情况.营业期限 }}',
    }
    for ri in range(len(t2.rows)):
        key = t2.cell(ri, 0).text.strip()
        if key in field_map_t2:
            set_cell_text(t2, ri, 1, field_map_t2[key])
        # Also handle "注册资本" which might have extra text
        elif '注册资本' in key:
            set_cell_text(t2, ri, 1, '{{ 基本情况.注册资本 | num }}')
    print("[OK] Table[2] 基本信息")

    # ============ TABLE[3]: 股东出资 (form.股东出资) ============
    t3 = doc.tables[3]
    # row0 = header, row1 = data (single xxxx row - need to convert to loop)
    # Rebuild as for / data / endfor
    for ci in range(len(t3.columns)):
        set_cell_text(t3, 1, ci, '')
    set_cell_text(t3, 1, 0, '{%tr for i in 股东出资 %}')
    set_cell_text(t3, 1, 1, '{{ i.认缴出资额 | money }}')
    set_cell_text(t3, 1, 2, '{{ i.认缴比例 | percent }}')
    set_cell_text(t3, 1, 3, '{{ i.实缴出资额 | money }}')
    if len(t3.columns) > 4:
        set_cell_text(t3, 1, 4, '{{ i.实缴比例 | percent }}')
    print("[OK] Table[3] 股东出资")

    # ============ TABLE[4]: 基本信息历史 (duplicate entity?) ============
    # This appears to be a second entity - keep as is for now
    # Just make it a duplicate of Table[2] pattern if needed
    t4 = doc.tables[4]
    for ri in range(len(t4.rows)):
        key = t4.cell(ri, 0).text.strip()
        if key in field_map_t2:
            set_cell_text(t4, ri, 1, field_map_t2[key])
    print("[OK] Table[4] 基本信息(第二实体)")

    # ============ TABLE[5]: 历史股东出资 (form.股东出资_历史) ============
    t5 = doc.tables[5]
    for ci in range(len(t5.columns)):
        set_cell_text(t5, 1, ci, '')
    set_cell_text(t5, 1, 0, '{%tr for i in 股东出资_历史 %}')
    set_cell_text(t5, 1, 1, '{{ i.认缴出资额 | money }}')
    set_cell_text(t5, 1, 2, '{{ i.认缴比例 | percent }}')
    set_cell_text(t5, 1, 3, '{{ i.实缴出资额 | money }}')
    if len(t5.columns) > 4:
        set_cell_text(t5, 1, 4, '{{ i.实缴比例 | percent }}')
    print("[OK] Table[5] 股东出资_历史")

    # ============ TABLE[6]: 审批手续清单 - KEEP AS-IS ============
    print("[OK] Table[6] 审批手续（保留原样）")

    # ============ TABLE[7]: PC合同分项造价 ============
    t7 = doc.tables[7]
    # row0 header, row1-4 data with xxxx
    cost_vars = ['设备费', '建筑安装工程费', '其他费', '合同总价']
    for ri, var_name in enumerate(cost_vars):
        r = ri + 1
        set_cell_text(t7, r, 0, str(r))
        set_cell_text(t7, r, 1, var_name)
        set_cell_text(t7, r, 3, '{{ PC合同.' + var_name + ' | money }}')
    print("[OK] Table[7] PC合同分项造价")

    # ============ TABLE[8]: 税率表 - KEEP AS-IS ============
    print("[OK] Table[8] 税率表（保留原样）")

    # ============ SPECIFIC PARAGRAPH REPLACEMENTS ============
    # Process all paragraphs for specific patterns
    for p in doc.paragraphs:
        full = ''.join(r.text for r in p.runs)

        # P[102]: 公司设立
        if '公司系由' in full and '出资组建' in full:
            for r in p.runs: r.text = ''
            new = '公司系由{{ 历史沿革.出资方 }}出资组建，于{{ 历史沿革.设立时间 }}取得{{ 历史沿革.核发机关 }}核发的、统一社会信用代码为{{ 基本情况.信用代码 }}的企业法人营业执照。设立时股权结构如下：'
            if p.runs: p.runs[0].text = new

        # P[108]: 项目概况
        if '项目位于' in full and '实际装机容量' in full:
            for r in p.runs: r.text = ''
            new = '{{ 项目概况.项目简称 }}项目位于{{ 项目概况.项目地址 }}，实际装机容量为{{ 项目概况.装机容量 | num }}MW，采用"自发自用，余量上网"运行方式。'
            if p.runs: p.runs[0].text = new

        # P[110]: 项目总投资
        if '竣工安装容量' in full and '不含税工程总价' in full:
            for r in p.runs: r.text = ''
            new = '{{ 项目概况.项目简称 }}项目竣工安装容量为{{ 项目概况.竣工容量 | num }}MW，不含税工程总价为{{ 项目概况.工程总价 | money }}元，不含税建设单价为{{ 项目概况.建设单价 | money }}元/Wp。'
            if p.runs: p.runs[0].text = new

        # P[114]: 承包单位
        if '乙方（承包单位）' in full and '【】' in full:
            for r in p.runs: r.text = ''
            new = '（1）乙方（承包单位）：{{ PC合同.承包单位 }}'
            if p.runs: p.runs[0].text = new

        # P[115]: 工程内容
        if '工程内容及规模' in full and '【】' in full:
            for r in p.runs: r.text = ''
            new = '（2）工程内容及规模：{{ PC合同.工程内容 }}'
            if p.runs: p.runs[0].text = new

        # P[118]: PC合同总价
        if '合同总价为' in full and '分项价格' in full:
            for r in p.runs: r.text = ''
            new = '（4）{{ 项目概况.项目简称 }}合同总价为{{ PC合同.合同总价 | money }}元，分项价格如下：'
            if p.runs: p.runs[0].text = new

        # P[120]: EMC合同名称
        if '合同名称' in full and '《xxxx》' in full:
            for r in p.runs: r.text = ''
            new = '①合同名称：《{{ EMC合同.合同名称 }}》'
            if p.runs: p.runs[0].text = new

        # P[121]: 用能方
        if '用能方' in full and '【】' in full:
            for r in p.runs: r.text = ''
            new = '②用能方：{{ EMC合同.用能方 }}'
            if p.runs: p.runs[0].text = new

        # P[122]: 房屋所有权证书登记人
        if '房屋所有权证书登记人' in full and '【】' in full:
            for r in p.runs: r.text = ''
            new = '③房屋所有权证书登记人：{{ EMC合同.登记人 }}'
            if p.runs: p.runs[0].text = new

        # P[126]: 结算方式 (empty / [])
        if '结算方式' in full and '[]' in full:
            for r in p.runs: r.text = ''
            new = '⑤结算方式：{{ EMC合同.结算方式 }}'
            if p.runs: p.runs[0].text = new

        # P[128]: 购售电合同
        if '截至尽调基准日' in full and '已并网发电' in full:
            for r in p.runs: r.text = ''
            new = '截至尽调基准日，{{ 项目概况.项目简称 }}项目已并网发电。根据与{{ 购售电合同.电网公司 }}签订的"非自然人分布式光伏发电项目购售电合同"，公司项目以余电上网模式消纳电量，电费由上网电费和补贴两部分组成，上网电费按上网电量与当地燃煤发电机组基准价（含脱硫、脱硝、除尘电价）乘积计算，补贴按发电量与对应的发电补贴标准（含税）乘积计算。'
            if p.runs: p.runs[0].text = new

    print("[OK] Paragraph replacements")

    # ============ TABLE CHECKS IN PARRANGE ============
    # Check for xxxx in tables and replace remaining
    for ti in [4]:
        t = doc.tables[ti]
        for ri in range(len(t.rows)):
            for ci in range(len(t.columns)):
                for p in t.cell(ri, ci).paragraphs:
                    ft = ''.join(r.text for r in p.runs)
                    if 'xxxx' in ft:
                        key = t.cell(ri, 0).text.strip()
                        if key in field_map_t2:
                            for r in p.runs: r.text = ''
                            if p.runs: p.runs[0].text = field_map_t2[key]

    # Save
    doc.save(str(OUT))
    print(f"\n[DONE] Saved: {OUT}")

    # Verify: count remaining 【】, xxxx, []
    doc2 = Document(str(OUT))
    all_text = ''
    for p in doc2.paragraphs:
        all_text += ''.join(r.text for r in p.runs) + '\n'
    for t in doc2.tables:
        for r in t.rows:
            for c in r.cells:
                all_text += c.text + '\n'

    remaining_placeholder = [c for c in ['【】', 'xxxx', '[]'] if c in all_text]
    jinja_count = all_text.count('{{')
    loop_count = all_text.count('{%tr')
    print(f"Remaining placeholders (【】/xxxx/[]): {remaining_placeholder}")
    print(f"Jinja2 {{ tags: {jinja_count}")
    print("{%tr for %} loops: " + str(loop_count))


if __name__ == '__main__':
    convert()
